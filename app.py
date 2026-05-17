"""
CALCULADORA DE VALOR DE RECONSTRUCCIÓN — SEGUROS CHILE
Conforme DFL 251, DS 1055, CCom art. 553, Ley 21.442 y NCG 556 CMF (dic. 2025)

Instalar: pip install streamlit pandas requests
Correr:   streamlit run app.py
Secrets (Streamlit Cloud): GOOGLE_MAPS_API_KEY = "AIza..."
"""

import streamlit as st
import pandas as pd
import requests
from datetime import date
from urllib.parse import quote
import re
import json
import subprocess
import tempfile
import os

# ─────────────────────────────────────────────────────────
# FORMATO NUMÉRICO CHILENO
# Separador de miles: punto (.)   Decimal: coma (,)
# Ejemplos: 6.000  /  227.500  /  1.250,50
# ─────────────────────────────────────────────────────────

def fmt_miles(n, decimales=0):
    """
    Formatea un número con punto como separador de miles y coma para decimales.
    Ej: fmt_miles(6000) → '6.000'
        fmt_miles(1234567.89, 2) → '1.234.567,89'
    """
    if n is None:
        return "—"
    if decimales == 0:
        s = f"{int(round(n)):,}".replace(",", ".")
    else:
        s = f"{n:,.{decimales}f}"
        # Python usa coma para miles y punto para decimal
        # Necesitamos invertir: '1,234,567.89' → '1.234.567,89'
        partes = s.split(".")
        entero = partes[0].replace(",", ".")
        decimal = partes[1] if len(partes) > 1 else ""
        s = f"{entero},{decimal}" if decimal else entero
    return s

def fmt_uf(n):
    """UF con 2 decimales. Ej: 227.500,00 UF"""
    return f"{fmt_miles(n, 2)} UF"

def fmt_m2(n):
    """m² sin decimales. Ej: 6.000 m²"""
    return f"{fmt_miles(n, 0)} m²"

def parsear_numero(texto):
    """
    Convierte texto ingresado por el usuario a float.
    Acepta formatos:
      - '6.000'      → 6000  (punto = miles)
      - '6000'       → 6000
      - '6,000'      → 6000  (coma = miles, formato anglosajón)
      - '1.250,50'   → 1250.50  (punto miles, coma decimal — formato chileno)
      - '1250.50'    → 1250.50  (punto decimal)
    Retorna float o None si no es válido.
    """
    if texto is None:
        return None
    t = str(texto).strip().replace(" ", "")
    if t == "" or t == "—":
        return None
    # Detectar formato chileno: tiene punto Y coma (1.250,50)
    if "." in t and "," in t:
        # Asumir punto=miles, coma=decimal
        t = t.replace(".", "").replace(",", ".")
    elif "," in t:
        # Solo coma: puede ser miles (6,000) o decimal (6,5)
        partes = t.split(",")
        if len(partes) == 2 and len(partes[1]) == 3 and partes[1].isdigit():
            # Es separador de miles anglosajón: 6,000
            t = t.replace(",", "")
        else:
            # Es decimal: 6,5
            t = t.replace(",", ".")
    elif "." in t:
        partes = t.split(".")
        if len(partes) == 2 and len(partes[1]) == 3 and partes[1].isdigit():
            # Es separador de miles: 6.000
            t = t.replace(".", "")
        # else: es decimal normal: 6.5 → se queda igual
    try:
        return float(t)
    except ValueError:
        return None

def input_numero(label, key, placeholder="", ayuda="", prefijo="", sufijo="",
                 es_entero=False, requerido=True):
    """
    Campo de texto que acepta números en formato chileno (punto=miles, coma=decimal).
    Retorna (valor_float_o_None, texto_ingresado, es_valido).
    Muestra el valor formateado en formato chileno bajo el campo.
    """
    col_in, col_fmt = st.columns([2, 1])
    with col_in:
        texto = st.text_input(
            label,
            key=key,
            placeholder=placeholder,
            help=ayuda or "Use punto para separar miles (ej: 6.000) y coma para decimales (ej: 1.250,50)",
        )
    valor = parsear_numero(texto)
    with col_fmt:
        st.markdown("<br>", unsafe_allow_html=True)
        if texto.strip() == "":
            st.caption(" ")
        elif valor is None:
            st.caption("⚠️ Valor inválido")
        else:
            if es_entero:
                st.caption(f"✓ {fmt_miles(valor, 0)} {sufijo}".strip())
            else:
                st.caption(f"✓ {fmt_miles(valor, 2)} {sufijo}".strip())
    es_valido = valor is not None
    return valor, texto, es_valido

# ─────────────────────────────────────────────────────────
# PARÁMETROS
# ─────────────────────────────────────────────────────────
FACTOR_GEOGRAFICO = {
    "Metropolitana (RM y ciudades grandes)": 1.05,
    "Intermedia (ciudades medianas)":        1.00,
    "Aislada (zonas rurales o extremas)":    1.15,
}
SISTEMAS_POR_TIPO = {
    "Casa":      ["Albañilería", "Metalcon"],
    "Depto":     ["Hormigón"],
    "Edificio":  ["Hormigón"],
    "Comunidad": ["Hormigón"],
}
NIVELES_POR_TS = {
    ("Casa","Albañilería"): ["Básico","Medio","Alto"],
    ("Casa","Metalcon"):    ["Básico","Medio"],
    ("Depto","Hormigón"):   ["Medio","Alto"],
    ("Edificio","Hormigón"):["Medio","Alto"],
    ("Comunidad","Hormigón"):["Medio","Alto"],
}
COSTOS_IND = {
    "Diseño del proyecto":      0.03,
    "Gastos generales de obra": 0.06,
    "Utilidad del contratista": 0.12,
    "Imprevistos":              0.10,
}
TASA_IVA = 0.19

ZONA_CORTA = {
    "Metropolitana (RM y ciudades grandes)": "Metropolitana",
    "Intermedia (ciudades medianas)":        "Intermedia",
    "Aislada (zonas rurales o extremas)":    "Aislada",
}
TS_LABEL = {
    ("Casa","Albañilería"):  "Casa / Albañilería",
    ("Casa","Metalcon"):     "Casa / Metalcon",
    ("Depto","Hormigón"):    "Depto / Hormigón",
    ("Edificio","Hormigón"): "Edificio / Hormigón",
    ("Comunidad","Hormigón"):"Comunidad / Hormigón",
}

# VUB referencias de mercado 2025-2026 (UF/m²)
REFS_VUB = {
    ("Metropolitana","Casa / Albañilería"):  {"Básico":(18,22),"Medio":(23,30),"Alto":(31,42)},
    ("Metropolitana","Casa / Metalcon"):     {"Básico":(16,20),"Medio":(21,28),"Alto":None},
    ("Metropolitana","Depto / Hormigón"):    {"Básico":None,   "Medio":(25,33),"Alto":(34,48)},
    ("Metropolitana","Edificio / Hormigón"): {"Básico":None,   "Medio":(26,35),"Alto":(36,52)},
    ("Metropolitana","Comunidad / Hormigón"):{"Básico":None,   "Medio":(25,34),"Alto":(35,50)},
    ("Intermedia","Casa / Albañilería"):     {"Básico":(17,21),"Medio":(22,29),"Alto":(30,40)},
    ("Intermedia","Casa / Metalcon"):        {"Básico":(15,19),"Medio":(20,27),"Alto":None},
    ("Intermedia","Depto / Hormigón"):       {"Básico":None,   "Medio":(24,32),"Alto":(33,46)},
    ("Intermedia","Edificio / Hormigón"):    {"Básico":None,   "Medio":(25,34),"Alto":(35,50)},
    ("Intermedia","Comunidad / Hormigón"):   {"Básico":None,   "Medio":(24,33),"Alto":(34,48)},
    ("Aislada","Casa / Albañilería"):        {"Básico":(20,26),"Medio":(27,36),"Alto":(37,50)},
    ("Aislada","Casa / Metalcon"):           {"Básico":(18,23),"Medio":(24,32),"Alto":None},
    ("Aislada","Depto / Hormigón"):          {"Básico":None,   "Medio":(29,38),"Alto":(39,55)},
    ("Aislada","Edificio / Hormigón"):       {"Básico":None,   "Medio":(30,40),"Alto":(41,58)},
    ("Aislada","Comunidad / Hormigón"):      {"Básico":None,   "Medio":(29,39),"Alto":(40,56)},
}

# ─────────────────────────────────────────────────────────
# MOTOR DE CÁLCULO
# ─────────────────────────────────────────────────────────
def factor_normativo(anio):
    if anio < 1985:  return 1.15
    if anio <= 2000: return 1.10
    if anio <= 2010: return 1.05
    return 1.00

def factor_altura(pisos):
    if pisos <= 2:  return 1.00
    if pisos <= 5:  return 1.05
    if pisos <= 10: return 1.10
    return 1.15

def calcular_vr(vub, sup, zona_label, pisos, anio, aplica_iva,
                fg_override=None, fn_override=None, fa_override=None,
                pct_diseno=None, pct_gg=None, pct_utilidad=None, pct_imprevistos=None):
    """
    Calcula el valor de reconstrucción.
    Acepta overrides de factores y porcentajes de costos indirectos.
    VUB se ingresa SIN IVA (la tabla MINVU y las referencias de mercado
    expresan costos directos netos; el IVA se aplica al final).
    """
    fg = fg_override if fg_override is not None else FACTOR_GEOGRAFICO[zona_label]
    fn = fn_override if fn_override is not None else factor_normativo(anio)
    fa = fa_override if fa_override is not None else factor_altura(pisos)

    # Porcentajes costos indirectos (editables o por defecto)
    p_dis = pct_diseno     if pct_diseno     is not None else COSTOS_IND["Diseño del proyecto"]
    p_gg  = pct_gg         if pct_gg         is not None else COSTOS_IND["Gastos generales de obra"]
    p_ut  = pct_utilidad   if pct_utilidad   is not None else COSTOS_IND["Utilidad del contratista"]
    p_imp = pct_imprevistos if pct_imprevistos is not None else COSTOS_IND["Imprevistos"]

    cd = sup * vub * fg * fn * fa
    ind_det = {
        "Diseño del proyecto":      cd * p_dis,
        "Gastos generales de obra": cd * p_gg,
        "Utilidad del contratista": cd * p_ut,
        "Imprevistos":              cd * p_imp,
    }
    ci  = sum(ind_det.values())
    st_ = cd + ci
    iv  = st_ * TASA_IVA if aplica_iva else 0.0
    return {
        "vub": vub, "fg": fg, "fn": fn, "fa": fa,
        "cd": cd, "ind_det": ind_det, "ci": ci, "st": st_,
        "iv": iv, "aplica_iva": aplica_iva, "vr": st_ + iv,
        "pcts": {"diseno": p_dis, "gg": p_gg, "utilidad": p_ut, "imprevistos": p_imp},
    }

def evaluar(monto, vr):
    if monto <= 0 or vr <= 0: return 0.0, False
    r = monto / vr
    return r, r < 1.0

def indemn(danio, monto, vr):
    ratio, infra = evaluar(monto, vr)
    return danio * ratio if infra else danio

# ─────────────────────────────────────────────────────────
# GOOGLE MAPS (opcional)
# ─────────────────────────────────────────────────────────
def get_gmaps_key():
    try:
        return st.secrets.get("GOOGLE_MAPS_API_KEY", "")
    except Exception:
        return ""

def geocodificar(direccion, api_key):
    try:
        r = requests.get(
            "https://maps.googleapis.com/maps/api/geocode/json",
            params={"address": direccion + ", Chile", "key": api_key},
            timeout=6,
        )
        data = r.json()
        if data.get("status") == "OK":
            loc = data["results"][0]["geometry"]["location"]
            return loc["lat"], loc["lng"], data["results"][0]["formatted_address"]
    except Exception:
        pass
    return None

# ─────────────────────────────────────────────────────────
# WIDGETS UI
# ─────────────────────────────────────────────────────────
def widget_vub(prefix, zona_label, tipo, sis, niv):
    """
    Muestra tabla de referencia VUB para la zona/tipo/nivel seleccionado
    y luego el campo de ingreso. La tabla aparece ANTES del campo.
    """
    zc  = ZONA_CORTA.get(zona_label, "")
    ts  = TS_LABEL.get((tipo, sis), "")
    ref = REFS_VUB.get((zc, ts), {})
    rng = ref.get(niv)

    # ── Tabla de referencia VUB (aparece ANTES del campo) ──
    filas_ref = [
        {"Nivel": nv, "Mín (UF/m²)": rg[0], "Máx (UF/m²)": rg[1],
         "Promedio ref.": round((rg[0]+rg[1])/2, 1),
         "Seleccionado": "✅" if nv == niv else ""}
        for nv, rg in (ref.items() if ref else {}.items())
        if rg
    ]

    with st.expander(
        f"📊 Tabla de referencia VUB — {tipo} / {sis} / zona {zc}",
        expanded=True,
    ):
        st.caption(
            "Rangos estimados de mercado 2025–2026. **No son valores oficiales.** "
            "Para el valor exacto consulte la tabla MINVU (en pesos, trimestral): "
            "[minvu.gob.cl](https://www.minvu.gob.cl/elementos-tecnicos/tabla-de-costos-unitarios/) "
            "o un tasador / corredor habilitado."
        )
        if filas_ref:
            st.dataframe(
                pd.DataFrame(filas_ref),
                use_container_width=True,
                hide_index=True,
            )
            if rng:
                prom = round((rng[0] + rng[1]) / 2, 1)
                st.info(
                    f"Para **{niv}** en zona **{zc}**: rango **{rng[0]}–{rng[1]} UF/m²** "
                    f"· promedio referencial **{prom} UF/m²**"
                )
        else:
            st.warning("Sin datos de referencia disponibles para esta combinación.")

    # ── Campo de ingreso VUB (aparece DESPUÉS de la tabla) ──
    ph = f"Ej: {round((rng[0]+rng[1])/2,1)} (rango {rng[0]}–{rng[1]})" if rng else "Ej: 28.0"
    return st.number_input(
        "VUB — Valor Unitario Base (UF/m²)",
        min_value=1.0, max_value=200.0, value=None,
        step=0.5, format="%.1f", placeholder=ph,
        key=f"{prefix}_vub",
        help="Ingrese el VUB según tasador, corredor de seguros o tabla MINVU convertida a UF.",
    )


def widget_tabla_refs_vub(zona_label):
    """Tabla completa de referencias VUB para la zona seleccionada."""
    zc = ZONA_CORTA.get(zona_label, "")
    filas = [
        {"Tipo / Sistema": ts, "Nivel": nv,
         "Mín (UF/m²)": rg[0], "Máx (UF/m²)": rg[1],
         "Promedio ref.": round((rg[0]+rg[1])/2, 1)}
        for (zc2, ts), nivs in REFS_VUB.items()
        if zc2 == zc
        for nv, rg in nivs.items()
        if rg
    ]
    if filas:
        with st.expander(f"📋 Ver tabla completa de referencias VUB — zona {zc}", expanded=False):
            st.caption("Estimaciones de mercado 2025–2026. No son valores oficiales. "
                       "Tabla MINVU (en pesos): https://www.minvu.gob.cl/elementos-tecnicos/tabla-de-costos-unitarios/")
            st.dataframe(
                pd.DataFrame(filas).sort_values(["Tipo / Sistema","Nivel"]),
                use_container_width=True, hide_index=True,
            )


def widget_distribucion_superficies(sup_total):
    """
    Panel visual completo de distribución bienes comunes / unidades privadas.
    Retorna (pct_comun_decimal, sup_comun_m2, sup_units_m2).
    """
    st.markdown("#### Distribución de superficies: bienes comunes vs unidades privadas")

    # ── Tabla de rangos referenciales ──
    st.markdown(
        "No existe un porcentaje único fijado por ley — cada edificio lo define en su "
        "**Reglamento de Copropiedad**. Use la tabla siguiente como referencia:"
    )

    tabla_html = """
<style>
.dist-table {width:100%;border-collapse:collapse;font-size:13px;margin-bottom:12px}
.dist-table th {background:#f0f2f6;padding:7px 10px;text-align:left;border-bottom:2px solid #ddd;font-weight:600}
.dist-table td {padding:6px 10px;border-bottom:1px solid #eee;vertical-align:top}
.dist-table tr:last-child td {background:#e8f4e8;font-weight:600}
.pct-comun {color:#1f77b4;font-weight:700;text-align:center}
.pct-units {color:#e67e22;font-weight:700;text-align:center}
.badge {display:inline-block;padding:2px 8px;border-radius:10px;font-size:11px;font-weight:600}
.b-blue {background:#dbeafe;color:#1d4ed8}
.b-orange {background:#fef3c7;color:#92400e}
.b-green {background:#d1fae5;color:#065f46}
</style>
<table class="dist-table">
  <tr>
    <th>Tipo de edificio</th>
    <th>Bienes comunes</th>
    <th>Unidades privadas</th>
    <th>Descripción</th>
  </tr>
  <tr>
    <td>🏢 Básico<br><small>2–5 pisos · sin amenidades</small></td>
    <td class="pct-comun"><span class="badge b-blue">25 – 35 %</span></td>
    <td class="pct-units"><span class="badge b-orange">65 – 75 %</span></td>
    <td><small>Pasillos, escaleras y conserjería básica. Sin subterráneos ni amenidades.</small></td>
  </tr>
  <tr>
    <td>🏢 Estándar<br><small>6–15 pisos · 1 subterráneo</small></td>
    <td class="pct-comun"><span class="badge b-blue">35 – 45 %</span></td>
    <td class="pct-units"><span class="badge b-orange">55 – 65 %</span></td>
    <td><small>Piscina, gimnasio, sala multiuso y 1 nivel de estacionamientos.</small></td>
  </tr>
  <tr>
    <td>🏢 Alto estándar<br><small>15–25 pisos · 2 subterráneos</small></td>
    <td class="pct-comun"><span class="badge b-blue">45 – 55 %</span></td>
    <td class="pct-units"><span class="badge b-orange">45 – 55 %</span></td>
    <td><small>Múltiples amenidades, lobby amplio, 2 niveles de estacionamientos.</small></td>
  </tr>
  <tr>
    <td>🏢 Premium / Torre<br><small>25+ pisos · 3+ subterráneos</small></td>
    <td class="pct-comun"><span class="badge b-blue">55 – 70 %</span></td>
    <td class="pct-units"><span class="badge b-orange">30 – 45 %</span></td>
    <td><small>Todas las amenidades, lobby doble altura, spa, múltiples subterráneos.</small></td>
  </tr>
  <tr>
    <td>📋 Referencia Ley 21.442</td>
    <td class="pct-comun"><span class="badge b-green">50 – 70 %</span><br><small>del VR total</small></td>
    <td class="pct-units">—</td>
    <td><small>Bienes comunes representan 50–70% del <strong>valor</strong> de reconstrucción (no solo superficie).</small></td>
  </tr>
</table>
<p style="font-size:11px;color:#888;margin-top:-8px">
Fuentes: Edifito / Ley 21.442 · ComunidadFeliz · OGUC art. 5.1.11 · Práctica de mercado 2025–2026
</p>
"""
    st.markdown(tabla_html, unsafe_allow_html=True)

    # ── Slider de selección ──
    pct_pct = st.slider(
        "Seleccione el % de superficie de bienes comunes para este edificio",
        min_value=20, max_value=70, value=40, step=1,
        help="Use la tabla anterior como guía. Para el valor exacto consulte el Reglamento de Copropiedad.",
        key="pct_comun_slider",
    )
    pct = pct_pct / 100

    # ── Barra visual proporcional ──
    bar_html = f"""
<div style="margin:8px 0 4px 0">
  <div style="display:flex;height:28px;border-radius:6px;overflow:hidden;border:1px solid #ddd">
    <div style="width:{pct_pct}%;background:#1f77b4;display:flex;align-items:center;
                justify-content:center;color:white;font-size:12px;font-weight:600;
                min-width:30px;transition:width 0.3s">
      {pct_pct}%
    </div>
    <div style="width:{100-pct_pct}%;background:#e67e22;display:flex;align-items:center;
                justify-content:center;color:white;font-size:12px;font-weight:600;
                min-width:30px;transition:width 0.3s">
      {100-pct_pct}%
    </div>
  </div>
  <div style="display:flex;gap:20px;margin-top:5px;font-size:12px">
    <span style="color:#1f77b4">■ Bienes comunes: <strong>{pct_pct}%</strong></span>
    <span style="color:#e67e22">■ Unidades privadas: <strong>{100-pct_pct}%</strong></span>
  </div>
</div>
"""
    st.markdown(bar_html, unsafe_allow_html=True)
    st.caption("⚠️ Distribución estimada. Para el valor exacto use el Reglamento de Copropiedad "
               "o los planos de arquitectura del edificio.")

    # ── Desglose en m² ──
    sup_comun = None
    sup_units = None
    if sup_total:
        sup_comun = round(sup_total * pct)
        sup_units = sup_total - sup_comun

        desglose_html = f"""
<div style="display:grid;grid-template-columns:1fr 1fr 1fr;gap:10px;margin:10px 0">
  <div style="background:#f8f9fa;border-radius:8px;padding:12px;border-left:4px solid #888;text-align:center">
    <div style="font-size:11px;color:#666;margin-bottom:4px">Superficie total</div>
    <div style="font-size:20px;font-weight:600">{fmt_miles(sup_total, 0)} m²</div>
  </div>
  <div style="background:#dbeafe;border-radius:8px;padding:12px;border-left:4px solid #1f77b4;text-align:center">
    <div style="font-size:11px;color:#1d4ed8;margin-bottom:4px">Bienes comunes ({pct_pct}%)</div>
    <div style="font-size:20px;font-weight:600;color:#1d4ed8">{fmt_miles(sup_comun, 0)} m²</div>
  </div>
  <div style="background:#fef3c7;border-radius:8px;padding:12px;border-left:4px solid #e67e22;text-align:center">
    <div style="font-size:11px;color:#92400e;margin-bottom:4px">Unidades privadas ({100-pct_pct}%)</div>
    <div style="font-size:20px;font-weight:600;color:#92400e">{fmt_miles(sup_units, 0)} m²</div>
  </div>
</div>
"""
        st.markdown(desglose_html, unsafe_allow_html=True)

    return pct, sup_comun, sup_units


def widget_herramienta_direccion(direccion):
    """
    Herramienta de apoyo para medir superficie desde dirección.
    Funciona con o sin API key de Google Maps.
    """
    if not direccion.strip():
        return

    with st.expander("🗺️ Herramienta de apoyo — Medir superficie desde dirección", expanded=False):
        st.markdown(
            "Use estas herramientas para estimar la **planta del edificio** "
            "y luego multiplíquela por el número de pisos + subterráneos."
        )

        api_key = get_gmaps_key()
        dir_encoded = quote(direccion + " Chile")

        # ── Links directos (siempre disponibles) ──
        col_ge, col_gm = st.columns(2)
        with col_ge:
            st.link_button(
                "🌍 Abrir en Google Earth",
                f"https://earth.google.com/web/search/{dir_encoded}",
                use_container_width=True,
            )
            st.caption("Herramienta Medir → Polígono → trace el contorno → obtenga el área en m²")
        with col_gm:
            st.link_button(
                "🗺️ Abrir en Google Maps",
                f"https://www.google.com/maps/search/{dir_encoded}",
                use_container_width=True,
            )
            st.caption("Clic derecho sobre el edificio → Medir distancia")

        st.divider()

        # ── Con API key: geocodificación automática ──
        if api_key:
            if st.button("📍 Geocodificar dirección y mostrar en mapa", key="geo_btn"):
                with st.spinner("Consultando Google Maps..."):
                    geo = geocodificar(direccion, api_key)
                if geo:
                    lat, lng, addr = geo
                    st.success(f"✅ Dirección encontrada: **{addr}**")
                    st.map(pd.DataFrame({"lat": [lat], "lon": [lng]}), zoom=17)
                    st.caption(f"Coordenadas: {lat:.6f}, {lng:.6f}")
                else:
                    st.error("No se encontró la dirección. Verifique el texto ingresado.")
        else:
            st.info(
                "**Búsqueda automática no activa** — agregue `GOOGLE_MAPS_API_KEY` "
                "en Secrets de Streamlit Cloud para activar el mapa integrado. "
                "Por ahora use los botones de arriba para abrir la dirección en Google Earth o Maps."
            )

        st.divider()
        st.markdown("**Pasos para medir en Google Earth:**")
        st.markdown(
            "1. Haga clic en el botón **Abrir en Google Earth** arriba\n"
            "2. El edificio aparecerá centrado en el mapa\n"
            "3. En el menú izquierdo, haga clic en **Medir** (ícono de regla)\n"
            "4. Seleccione **Polígono** y trace el contorno exterior del edificio\n"
            "5. Google Earth muestra el **área en m²** — ese es el área de la planta\n"
            "6. Multiplique: **área planta × (N° pisos + N° subterráneos)** = superficie total\n"
            "7. Ingrese ese valor en el campo de superficie del formulario"
        )


def widget_formulario_componente(prefix, zona, pisos, anio, aplica_iva,
                                  default_tipo="Comunidad", label_tipo="Tipo de inmueble",
                                  sup_sugerida=None, mostrar_dist=False,
                                  pct_comun=None, pct_pct=None,
                                  sup_total=None):
    """
    Formulario completo de un componente.
    Si mostrar_dist=True, muestra la tabla de distribución de superficies
    justo debajo del campo Superficie (m²).
    """
    tipos = list(SISTEMAS_POR_TIPO.keys())
    idx   = tipos.index(default_tipo) if default_tipo in tipos else 0
    tipo  = st.selectbox(label_tipo, tipos, index=idx, key=f"{prefix}_tipo")
    sis   = st.selectbox("Sistema constructivo", SISTEMAS_POR_TIPO[tipo], key=f"{prefix}_sis")
    niv   = st.selectbox("Nivel de terminaciones", NIVELES_POR_TS[(tipo, sis)],
                         key=f"{prefix}_niv",
                         help="Básico = sin lujos · Medio = estándar · Alto = premium")

    # ── Tabla referencia VUB (antes del campo VUB) ──
    if zona:
        vub = widget_vub(prefix, zona, tipo, sis, niv)
    else:
        vub = st.number_input("VUB — Valor Unitario Base (UF/m²)", min_value=1.0,
                              max_value=200.0, value=None, step=0.5, format="%.1f",
                              key=f"{prefix}_vub")

    st.caption(
        "ℹ️ **El VUB se ingresa sin IVA.** Las tablas de referencia (MINVU y mercado privado) "
        "expresan los costos de construcción en valores netos. El IVA (19%) se aplica al final."
    )

    # ── Factores siempre visibles (sin expander) ──
    st.markdown("**Factores del cálculo**")
    st.caption("Calculados automáticamente según zona, año y pisos. Ajuste solo si dispone de datos más precisos.")
    fa1, fa2, fa3 = st.columns(3)
    with fa1:
        fg_edit = st.number_input(
            "Factor geográfico",
            min_value=0.50, max_value=2.00,
            value=round(float(FACTOR_GEOGRAFICO.get(zona, 1.0)) if zona else 1.0, 2),
            step=0.01, format="%.2f", key=f"{prefix}_fg",
            help="Metropolitana 1,05 · Intermedia 1,00 · Aislada 1,15",
        )
    with fa2:
        fn_default = round(factor_normativo(anio), 2) if anio else 1.0
        fn_edit = st.number_input(
            f"Factor normativo (año {anio or '—'})",
            min_value=0.50, max_value=2.00,
            value=fn_default,
            step=0.01, format="%.2f", key=f"{prefix}_fn",
            help="<1985: 1,15 · 1985-2000: 1,10 · 2001-2010: 1,05 · >2010: 1,00",
        )
    with fa3:
        fa_default = round(factor_altura(pisos), 2) if pisos else 1.0
        fa_edit = st.number_input(
            f"Factor altura ({pisos or '—'} pisos)",
            min_value=0.50, max_value=2.00,
            value=fa_default,
            step=0.01, format="%.2f", key=f"{prefix}_fa",
            help="1-2 pisos: 1,00 · 3-5: 1,05 · 6-10: 1,10 · 11+: 1,15",
        )

    st.markdown("**Costos indirectos (% sobre costo directo)**")
    fp1, fp2, fp3, fp4 = st.columns(4)
    with fp1:
        p_dis = st.number_input("Diseño (%)", min_value=0.0, max_value=30.0,
                                value=3.0, step=0.5, format="%.1f",
                                key=f"{prefix}_pdis") / 100
    with fp2:
        p_gg  = st.number_input("Gastos grales. (%)", min_value=0.0, max_value=30.0,
                                value=6.0, step=0.5, format="%.1f",
                                key=f"{prefix}_pgg") / 100
    with fp3:
        p_ut  = st.number_input("Utilidad contr. (%)", min_value=0.0, max_value=30.0,
                                value=12.0, step=0.5, format="%.1f",
                                key=f"{prefix}_put") / 100
    with fp4:
        p_imp = st.number_input("Imprevistos (%)", min_value=0.0, max_value=30.0,
                                value=10.0, step=0.5, format="%.1f",
                                key=f"{prefix}_pimp") / 100
    total_ind = (p_dis + p_gg + p_ut + p_imp) * 100
    st.caption(f"Total costos indirectos: **{total_ind:.1f}%** sobre el costo directo.")

    # ── Campo Superficie (m²) ──
    sup_val, sup_txt, sup_ok = input_numero(
        "Superficie (m²)", key=f"{prefix}_sup",
        placeholder="Ej: 3.500",
        ayuda="Ingrese la superficie con punto como separador de miles. Ej: 3.500 o 12.000",
        sufijo="m²", es_entero=True,
    )
    sup = int(sup_val) if sup_val else (sup_sugerida if sup_sugerida else None)

    # ── Vista previa del cálculo en tiempo real ──
    if vub and sup and fg_edit and fn_edit and fa_edit:
        cd_p = sup * vub * fg_edit * fn_edit * fa_edit
        ci_p = cd_p * (p_dis + p_gg + p_ut + p_imp)
        st_p = cd_p + ci_p
        iv_p = st_p * TASA_IVA if aplica_iva else 0
        vr_p = st_p + iv_p
        st.info(
            f"**Vista previa:** {fmt_m2(sup)} × {vub:.1f} UF/m² "
            f"× {fg_edit:.2f} × {fn_edit:.2f} × {fa_edit:.2f} "
            f"= **{fmt_uf(cd_p)}** (CD) "
            f"→ +{total_ind:.0f}% ind."
            + (f" +IVA 19%" if aplica_iva else "")
            + f" = **{fmt_uf(vr_p)}**"
        )

    # ── Panel de distribución (aparece siempre debajo de Superficie cuando mostrar_dist=True) ──
    if mostrar_dist and pct_comun is not None and pct_pct is not None:
        # Usa la superficie ingresada en este campo si no viene sup_total
        sup_ref = sup_total if sup_total else sup
        if sup_ref:
            sup_comun_ref = round(sup_ref * pct_comun)
            sup_units_ref = sup_ref - sup_comun_ref
            dist_html = f"""
<div style="margin:8px 0 14px 0;padding:12px 14px;background:#f8faff;
            border:1px solid #c7d8f0;border-left:4px solid #1f77b4;
            border-radius:0 6px 6px 0;font-size:13px">
  <div style="font-weight:600;margin-bottom:8px;color:#1d3557">
    📐 Distribución de superficies según porcentaje seleccionado
  </div>
  <div style="display:grid;grid-template-columns:1fr 1fr 1fr;gap:8px;margin-bottom:8px">
    <div style="background:white;border-radius:6px;padding:8px 10px;text-align:center;
                border:1px solid #ddd">
      <div style="font-size:11px;color:#555;margin-bottom:3px">Superficie ingresada</div>
      <div style="font-size:18px;font-weight:700;color:#333">{fmt_miles(sup_ref, 0)} m²</div>
    </div>
    <div style="background:#dbeafe;border-radius:6px;padding:8px 10px;text-align:center;
                border:1px solid #93c5fd">
      <div style="font-size:11px;color:#1d4ed8;margin-bottom:3px">■ Bienes comunes ({pct_pct}%)</div>
      <div style="font-size:18px;font-weight:700;color:#1d4ed8">{fmt_miles(sup_comun_ref, 0)} m²</div>
    </div>
    <div style="background:#fef3c7;border-radius:6px;padding:8px 10px;text-align:center;
                border:1px solid #fcd34d">
      <div style="font-size:11px;color:#92400e;margin-bottom:3px">■ Unidades privadas ({100-pct_pct}%)</div>
      <div style="font-size:18px;font-weight:700;color:#92400e">{fmt_miles(sup_units_ref, 0)} m²</div>
    </div>
  </div>
  <div style="font-size:11px;color:#666">
    ⚠️ Distribución estimada según el porcentaje seleccionado ({pct_pct}% / {100-pct_pct}%).
    Para el valor exacto consulte el Reglamento de Copropiedad o los planos del edificio.
  </div>
</div>
"""
            st.markdown(dist_html, unsafe_allow_html=True)
        else:
            # Superficie aún no ingresada — mostrar aviso
            st.caption(
                "💡 Una vez que ingrese la superficie, se mostrará el desglose estimado "
                f"entre bienes comunes ({pct_pct}%) y unidades privadas ({100-pct_pct}%)."
            )

    # ── Monto asegurado ──
    monto_val, monto_txt, monto_ok = input_numero(
        "Monto asegurado en póliza (UF)", key=f"{prefix}_monto",
        placeholder="Ej: 227.500",
        ayuda="Ingrese el monto con punto como separador de miles. Ej: 227.500. Ingrese 0 si no hay seguro.",
        sufijo="UF",
    )
    monto = monto_val if monto_val is not None else 0

    return {"tipo":tipo,"sis":sis,"niv":niv,"vub":vub,"sup":sup,"monto":monto,
            "zona":zona,"pisos":pisos,"anio":anio,"aplica_iva":aplica_iva,
            "fg_edit":fg_edit,"fn_edit":fn_edit,"fa_edit":fa_edit,
            "p_dis":p_dis,"p_gg":p_gg,"p_ut":p_ut,"p_imp":p_imp}


def validar_comp(d, campo):
    errs = []
    if not d.get("vub"):
        errs.append(f"Ingrese el VUB (UF/m²) de {campo}.")
    if not d.get("sup"):
        errs.append(f"Ingrese la superficie de {campo} (ej: 3.500 m²).")
    if d.get("monto") is None:
        errs.append(f"Ingrese el monto asegurado de {campo} (puede ser 0).")
    return errs


def widget_resultado(label, res, datos, danio_pct, perdida_real=None, nota="", expanded=True):
    """Muestra resultado de un componente calculado."""
    vr    = res["vr"]
    monto = datos.get("monto") or 0
    ratio, infra = evaluar(monto, vr)
    # Calcular daño: prioridad a pérdida real ingresada, luego % simulado
    if perdida_real and perdida_real > 0:
        danio_ = perdida_real
        origen_danio = f"pérdida real ingresada ({fmt_uf(perdida_real)})"
    elif danio_pct and danio_pct > 0:
        danio_ = vr * (danio_pct / 100)
        origen_danio = f"simulación {danio_pct}% del VR"
    else:
        danio_ = None
        origen_danio = None
    ind_ = indemn(danio_, monto, vr) if danio_ is not None else None
    pcts = res.get("pcts", {})

    with st.expander(f"📦 {label} — **{fmt_uf(vr)}**", expanded=expanded):
        if nota:
            st.caption(nota)
        if monto <= 0:
            st.info("ℹ️ Sin monto asegurado registrado.")
        elif infra:
            st.warning(
                f"⚠️ **Infrasegurado.** Cobertura: **{ratio*100:.1f}%** "
                f"— Brecha: **{fmt_uf(vr - monto)}**"
            )
        else:
            st.success(f"✅ Cobertura adecuada ({ratio*100:.1f}%)")

        c1, c2, c3 = st.columns(3)
        c1.metric("Valor de reconstrucción", fmt_uf(vr))
        c2.metric("Monto asegurado", fmt_uf(monto) if monto > 0 else "No indicado")
        c3.metric("Cobertura", f"{ratio*100:.1f}%" if monto > 0 else "—",
                  delta=f"{(ratio-1)*100:.1f}%" if monto > 0 else None,
                  delta_color="normal" if not infra else "inverse")
        if monto > 0:
            st.progress(min(ratio, 1.0), text=f"Cobertura: {ratio*100:.1f}%")

        st.markdown("**Desglose del cálculo**")
        sup_disp = datos.get("sup") or 0
        st.markdown(f"""
| # | Concepto | Valor |
|---|----------|-------|
| 1 | VUB ingresado *(sin IVA)* — {datos.get('tipo','')}/{datos.get('sis','')}/{datos.get('niv','')} | **{fmt_miles(res['vub'],1)} UF/m²** |
| 2 | × Factor geográfico | {res['fg']:.2f} |
| 3 | × Factor normativo (año {datos.get('anio','')}) | {res['fn']:.2f} |
| 4 | × Factor altura ({datos.get('pisos','')} pisos) | {res['fa']:.2f} |
| 5 | **Costo directo** ({fmt_m2(sup_disp)}) | **{fmt_uf(res['cd'])}** |
| 6a | + Diseño del proyecto ({pcts.get('diseno',0.03)*100:.1f}%) | {fmt_uf(res['ind_det']['Diseño del proyecto'])} |
| 6b | + Gastos generales de obra ({pcts.get('gg',0.06)*100:.1f}%) | {fmt_uf(res['ind_det']['Gastos generales de obra'])} |
| 6c | + Utilidad del contratista ({pcts.get('utilidad',0.12)*100:.1f}%) | {fmt_uf(res['ind_det']['Utilidad del contratista'])} |
| 6d | + Imprevistos ({pcts.get('imprevistos',0.10)*100:.1f}%) | {fmt_uf(res['ind_det']['Imprevistos'])} |
| 7 | **Subtotal sin IVA** | **{fmt_uf(res['st'])}** |
| 8 | + IVA 19% *(sobre subtotal neto)* | {fmt_uf(res['iv'])} |
| ✓ | **VALOR DE RECONSTRUCCIÓN** | **{fmt_uf(vr)}** |
""")

        # Simulación: solo si hay daño (% o pérdida real)
        if danio_ is not None and danio_ > 0:
            st.markdown(f"**Simulación de siniestro** — {origen_danio}")
            s1, s2, s3 = st.columns(3)
            s1.metric("Daño / Pérdida", fmt_uf(danio_))
            s2.metric("Indemnización real", fmt_uf(ind_) if monto > 0 else "—")
            if infra:
                s3.metric("Pérdida no cubierta", fmt_uf(danio_ - ind_), delta_color="inverse")
                st.warning(
                    f"**Art. 553 CCom:** recibiría **{fmt_uf(ind_)}** en vez de **{fmt_uf(danio_)}**. "
                    f"Pérdida no cubierta: **{fmt_uf(danio_ - ind_)}**."
                )


# ─────────────────────────────────────────────────────────
# INFORME TXT
# ─────────────────────────────────────────────────────────
def _bloque_txt(etiq, res, datos, danio_pct):
    monto = datos.get("monto") or 0
    vr    = res["vr"]
    ratio, infra = evaluar(monto, vr)
    d = vr * (danio_pct / 100)
    i = indemn(d, monto, vr)
    lns = [
        f"  [{etiq}]",
        f"    Tipo/Sistema/Nivel  : {datos.get('tipo','')}/{datos.get('sis','')}/{datos.get('niv','')}",
        f"    VUB ingresado       : {res['vub']:.1f} UF/m²",
        f"    Superficie          : {datos.get('sup',0):,.0f} m²",
        f"    Factores            : geográfico {res['fg']:.2f} · normativo {res['fn']:.2f} · altura {res['fa']:.2f}",
        f"    Costo directo       : {res['cd']:>12,.2f} UF",
        f"    Costos ind. (31%)   : {res['ci']:>12,.2f} UF",
        f"    Subtotal s/IVA      : {res['st']:>12,.2f} UF",
        (f"    IVA 19%             : {res['iv']:>12,.2f} UF" if res["aplica_iva"] else
         f"    IVA 19%             :       no aplica"),
        f"    VALOR RECONSTRUCCIÓN: {vr:>12,.2f} UF",
        "",
        (f"    Monto asegurado     : {monto:>12,.2f} UF" if monto > 0 else
         f"    Monto asegurado     :    No indicado"),
        (f"    Cobertura           : {ratio*100:>11.1f} %" if monto > 0 else
         f"    Cobertura           :            —"),
        f"    Infraseguro         : {'SÍ ⚠' if infra else 'NO ✓'}",
    ]
    if infra:
        lns.append(f"    Brecha sin cubrir   : {vr-monto:>12,.2f} UF")
    lns += [
        f"    Simulación ({danio_pct:.0f}% daño)",
        f"      Daño estimado     : {d:>12,.2f} UF",
        (f"      Indemnización     : {i:>12,.2f} UF" if monto > 0 else
         f"      Indemnización     :    Ver VR"),
    ]
    if infra:
        lns.append(f"      Pérdida           : {d-i:>12,.2f} UF")
    return "\n".join(lns)


def generar_word(caso, perdida_real=None):
    """
    Genera un archivo .docx usando python-docx (sin Node.js).
    Retorna bytes del archivo o None si falla.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io

    def rgb(hex_str):
        h = hex_str.lstrip('#')
        return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

    def set_cell_bg(cell, color_hex):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex.lstrip('#'))
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    def add_heading(doc, text, level=1):
        p = doc.add_paragraph()
        p.style = f'Heading {level}'
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14 if level==1 else 12)
        run.font.color.rgb = rgb('1F3864')
        run.font.bold = True
        return p

    def add_para(doc, text, bold=False, size=10, color='333333', space_after=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = rgb(color)
        return p

    def add_table_info(doc, rows):
        """Tabla de 2 columnas: etiqueta / valor"""
        t = doc.add_table(rows=len(rows), cols=2)
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, (k, v) in enumerate(rows):
            c0, c1 = t.rows[i].cells
            c0.width = Cm(5.5); c1.width = Cm(10.5)
            c0.paragraphs[0].clear()
            r0 = c0.paragraphs[0].add_run(str(k))
            r0.font.name='Arial'; r0.font.size=Pt(9); r0.font.bold=True
            c0.paragraphs[0].paragraph_format.space_after = Pt(2)
            c1.paragraphs[0].clear()
            r1 = c1.paragraphs[0].add_run(str(v))
            r1.font.name='Arial'; r1.font.size=Pt(9)
            c1.paragraphs[0].paragraph_format.space_after = Pt(2)
            set_cell_bg(c0, 'EEF3FB' if i%2==0 else 'FFFFFF')
            set_cell_bg(c1, 'EEF3FB' if i%2==0 else 'FFFFFF')
        return t

    def add_table_desglose(doc, filas):
        """Tabla de desglose del cálculo: #, concepto, valor"""
        headers = ['#', 'Concepto', 'Valor']
        widths  = [Cm(1), Cm(10), Cm(5)]
        t = doc.add_table(rows=1+len(filas), cols=3)
        t.style = 'Table Grid'
        # Header
        for j,(h,w) in enumerate(zip(headers, widths)):
            cell = t.rows[0].cells[j]
            cell.width = w
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(h)
            run.font.name='Arial'; run.font.size=Pt(9); run.font.bold=True; run.font.color.rgb=rgb('FFFFFF')
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            set_cell_bg(cell, '1F3864')
        # Filas
        for i, (num, conc, val) in enumerate(filas):
            row = t.rows[i+1]
            shade = 'F5F8FF' if i%2==0 else 'FFFFFF'
            row.cells[0].width=widths[0]; row.cells[1].width=widths[1]; row.cells[2].width=widths[2]
            for j,(txt,aln) in enumerate([(str(num),'center'),(str(conc),'left'),(str(val),'right')]):
                cell = row.cells[j]
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(txt)
                run.font.name='Arial'; run.font.size=Pt(9)
                if str(num)=='✓': run.font.bold=True
                cell.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.CENTER if aln=='center'
                                                 else WD_ALIGN_PARAGRAPH.RIGHT if aln=='right'
                                                 else WD_ALIGN_PARAGRAPH.LEFT)
                cell.paragraphs[0].paragraph_format.space_after = Pt(2)
                set_cell_bg(cell, shade)
        return t

    def bloque_componente(doc, etiq, comp, dp, pr):
        res   = comp.get('res', {})
        vr    = res.get('vr', 0)
        monto = comp.get('monto') or 0
        ratio = monto/vr if vr>0 and monto>0 else 0
        infra = monto>0 and ratio<1.0
        pcts  = res.get('pcts', {})

        # Daño
        if pr and pr>0:
            danio=pr; origen=f"Pérdida real ({fmt_uf(pr)})"
        elif dp and dp>0:
            danio=vr*dp/100; origen=f"Simulación {dp}%"
        else:
            danio=None; origen=None
        indemn_v = (danio*ratio if infra else danio) if danio else None

        add_heading(doc, etiq, level=2)

        filas = [
            ('1', f"VUB ingresado (sin IVA) — {comp.get('tipo','')}/{comp.get('sis','')}/{comp.get('niv','')}", f"{fmt_miles(res.get('vub',0),1)} UF/m²"),
            ('2', '× Factor geográfico', f"{res.get('fg',0):.2f}"),
            ('3', f"× Factor normativo (año {comp.get('anio','')})", f"{res.get('fn',0):.2f}"),
            ('4', f"× Factor altura ({comp.get('pisos','')} pisos)", f"{res.get('fa',0):.2f}"),
            ('5', f"Costo directo ({fmt_m2(comp.get('sup',0))})", fmt_uf(res.get('cd',0))),
            ('6a', f"+ Diseño del proyecto ({(pcts.get('diseno',0.03)*100):.1f}%)", fmt_uf(res.get('ind_det',{}).get('Diseño del proyecto',0))),
            ('6b', f"+ Gastos generales de obra ({(pcts.get('gg',0.06)*100):.1f}%)", fmt_uf(res.get('ind_det',{}).get('Gastos generales de obra',0))),
            ('6c', f"+ Utilidad del contratista ({(pcts.get('utilidad',0.12)*100):.1f}%)", fmt_uf(res.get('ind_det',{}).get('Utilidad del contratista',0))),
            ('6d', f"+ Imprevistos ({(pcts.get('imprevistos',0.10)*100):.1f}%)", fmt_uf(res.get('ind_det',{}).get('Imprevistos',0))),
            ('7', 'Subtotal sin IVA', fmt_uf(res.get('st',0))),
            ('8', '+ IVA 19% (sobre subtotal neto)', fmt_uf(res.get('iv',0))),
            ('✓', 'VALOR DE RECONSTRUCCIÓN', fmt_uf(vr)),
        ]
        add_table_desglose(doc, filas)
        doc.add_paragraph()

        # Análisis póliza
        add_para(doc, 'Análisis de póliza', bold=True, color='1F3864')
        filas_pol = [
            ('Valor de reconstrucción', fmt_uf(vr)),
            ('Monto asegurado en póliza', fmt_uf(monto) if monto>0 else 'No indicado'),
            ('Cobertura', f"{ratio*100:.1f}%" if monto>0 else '—'),
            ('Infraseguro', 'SÍ ⚠' if infra else ('NO ✓' if monto>0 else '—')),
        ]
        if infra: filas_pol.append(('Brecha sin cubrir', fmt_uf(vr-monto)))
        add_table_info(doc, filas_pol)
        doc.add_paragraph()

        # Simulación
        if danio and danio>0:
            add_para(doc, f'Simulación de siniestro — {origen}', bold=True, color='1F3864')
            filas_sim = [
                ('Daño / Pérdida', fmt_uf(danio)),
                ('Indemnización real', fmt_uf(indemn_v) if monto>0 else 'Ver VR'),
            ]
            if infra and monto>0: filas_sim.append(('Pérdida no cubierta', fmt_uf(danio-indemn_v)))
            add_table_info(doc, filas_sim)
            doc.add_paragraph()

    try:
        doc = Document()

        # Configurar página A4
        section = doc.sections[0]
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin  = section.bottom_margin = Cm(2.0)

        # Estilos base
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

        dp = caso.get('danio_pct', 0)
        pr = perdida_real or 0

        # ── Título ──
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('INFORME DE VALOR DE RECONSTRUCCIÓN')
        run.font.name='Arial'; run.font.size=Pt(16); run.font.bold=True
        run.font.color.rgb = rgb('1F3864')
        p.paragraph_format.space_after = Pt(4)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run('Conforme DFL 251 · DS 1055 · CCom art. 553 · Ley 21.442 · NCG 556 CMF')
        r2.font.name='Arial'; r2.font.size=Pt(9); r2.font.color.rgb=rgb('666666')
        doc.add_paragraph()

        # ── Identificación ──
        add_heading(doc, '1. Identificación de la propiedad', level=1)
        add_table_info(doc, [
            ('Nombre / Referencia', caso.get('nombre','—')),
            ('Dirección',           caso.get('direccion','—')),
            ('Zona geográfica',     caso.get('zona','—')),
            ('Número de pisos',     str(caso.get('pisos','—'))),
            ('Año de construcción', str(caso.get('anio','—'))),
            ('Fecha de cálculo',    date.today().strftime('%d/%m/%Y')),
        ])
        doc.add_paragraph()

        # ── Contenido según modo ──
        modo = caso.get('modo', 'simple')
        if modo == 'simple':
            add_heading(doc, '2. Cálculo — Inmueble completo', level=1)
            bloque_componente(doc, 'Inmueble completo', caso.get('comp',{}), dp, pr)

        elif modo == 'comunes':
            add_heading(doc, '2. Bienes y espacios comunes (Ley 21.442 art. 43)', level=1)
            comp = caso.get('comp', {})
            if comp.get('sup_total'):
                add_para(doc,
                    f"Superficie total: {fmt_m2(comp.get('sup_total',0))} · "
                    f"Bienes comunes ({round(comp.get('pct_comun',0)*100):.0f}%): {fmt_m2(comp.get('sup',0))}",
                    bold=True, color='1F3864')
            bloque_componente(doc, 'Bienes y espacios comunes', comp, dp, pr)

        elif modo == 'comunidad':
            add_heading(doc, '2. Póliza colectiva — NCG 556 CMF', level=1)
            add_heading(doc, 'Bloque 1 — Bienes y espacios comunes', level=2)
            add_para(doc, 'Asegurado: la comunidad / condominio. OBLIGATORIO (Ley 21.442 art. 43).', bold=True, color='1F3864')
            bloque_componente(doc, 'Bienes comunes', caso.get('comp_comun',{}), dp, pr)

            add_heading(doc, 'Bloque 2 — Unidades privadas', level=2)
            add_para(doc, 'Asegurado: cada copropietario individualmente.', color='555555')
            for u in caso.get('unidades', []):
                bloque_componente(doc, u.get('nombre','Unidad'), u, dp, pr)

            # Consolidado
            vr_t = caso.get('total_vr', 0)
            m_t  = caso.get('total_monto', 0)
            r_t  = m_t/vr_t if vr_t>0 and m_t>0 else 0
            i_t  = m_t>0 and r_t<1
            if pr>0: d_t=pr; orig_t=f"Pérdida real ({fmt_uf(pr)})"
            elif dp>0: d_t=vr_t*dp/100; orig_t=f"Simulación {dp}%"
            else: d_t=None; orig_t=None
            ind_t = (d_t*r_t if i_t else d_t) if d_t else None

            add_heading(doc, '3. Consolidado total de la comunidad', level=1)
            filas_cons = [
                ('VR bienes comunes',    fmt_uf(caso.get('vr_comun',0))),
                ('VR unidades privadas', fmt_uf(caso.get('vr_units',0))),
                ('VR TOTAL COMUNIDAD',   fmt_uf(vr_t)),
                ('Monto asegurado total',fmt_uf(m_t) if m_t>0 else 'No indicado'),
                ('Cobertura global',     f"{r_t*100:.1f}%" if m_t>0 else '—'),
                ('Infraseguro global',   'SÍ ⚠' if i_t else ('NO ✓' if m_t>0 else '—')),
            ]
            if i_t: filas_cons.append(('Brecha total', fmt_uf(vr_t-m_t)))
            if d_t: filas_cons.append(('Simulación / Pérdida', f"{orig_t}: {fmt_uf(d_t)}"))
            if d_t and m_t>0: filas_cons.append(('Indemnización global', fmt_uf(ind_t)))
            if d_t and i_t and m_t>0: filas_cons.append(('Pérdida no cubierta', fmt_uf(d_t-ind_t)))
            add_table_info(doc, filas_cons)

        # ── Pie legal ──
        doc.add_paragraph()
        p_pie = doc.add_paragraph()
        p_pie.paragraph_format.space_before = Pt(12)
        run_pie = p_pie.add_run(
            "Nota: Informe referencial. Los valores deben verificarse con un tasador habilitado. "
            "VUB ingresado por el usuario. Tabla MINVU: www.minvu.gob.cl · "
            "Normativa: DFL 251 · DS 1055 · CCom 553 · Ley 21.442 · NCG 556 CMF"
        )
        run_pie.font.name='Arial'; run_pie.font.size=Pt(8); run_pie.font.italic=True
        run_pie.font.color.rgb=rgb('888888')

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    except Exception as e:
        st.error(f"Error generando Word: {e}")
        return None


def generar_informe(caso):
    sep  = "=" * 64
    sep2 = "─" * 64
    hoy  = date.today().strftime("%d/%m/%Y")
    lns  = [
        "INFORME DE VALOR DE RECONSTRUCCIÓN",
        "Conforme DFL 251, DS 1055, CCom art. 553 y Ley 21.442",
        sep,
        f"  Nombre / Referencia : {caso.get('nombre','—')}",
        f"  Dirección           : {caso.get('direccion','—')}",
        f"  Zona geográfica     : {caso.get('zona','—')}",
        f"  Número de pisos     : {caso.get('pisos','—')}",
        f"  Año de construcción : {caso.get('anio','—')}",
        f"  Fecha de cálculo    : {hoy}",
        sep, "",
    ]
    modo = caso.get("modo","simple")
    if modo == "simple":
        lns += ["INMUEBLE COMPLETO\n",
                _bloque_txt("Inmueble completo", caso["comp"]["res"], caso["comp"], caso["danio_pct"])]
    elif modo == "comunes":
        lns += ["BIENES Y ESPACIOS COMUNES (Ley 21.442 art. 43)\n",
                _bloque_txt("Bienes comunes", caso["comp"]["res"], caso["comp"], caso["danio_pct"])]
    elif modo == "comunidad":
        if caso.get("desglose"):
            dg = caso["desglose"]
            lns += [
                "DISTRIBUCIÓN DE SUPERFICIES",
                f"  Superficie total         : {dg.get('sup_total',0):,.0f} m²",
                f"  % bienes comunes         : {dg.get('pct_pct',0)}%",
                f"  Superficie bienes comunes: {dg.get('sup_comun',0):,.0f} m²",
                f"  Superficie unidades      : {dg.get('sup_units',0):,.0f} m²",
                f"  Referencia               : Tabla Ley 21.442 / OGUC / mercado 2025",
                "",
            ]
        lns += ["PÓLIZA COLECTIVA — NCG 556 CMF", "",
                "BLOQUE 1: BIENES Y ESPACIOS COMUNES (asegurado: la comunidad)\n",
                _bloque_txt("Bienes comunes", caso["comp_comun"]["res"],
                            caso["comp_comun"], caso["danio_pct"]),
                "", sep2, "",
                "BLOQUE 2: UNIDADES PRIVADAS (asegurado: cada copropietario)\n"]
        for u in caso.get("unidades", []):
            lns += [_bloque_txt(u.get("nombre") or "Unidad", u["res"], u, caso["danio_pct"]), ""]
        vr_t = caso.get("total_vr", 0)
        m_t  = caso.get("total_monto", 0)
        r_t, i_t = evaluar(m_t, vr_t)
        d_t  = vr_t * (caso.get("danio_pct", 0) / 100)
        ind_t = indemn(d_t, m_t, vr_t)
        lns += [sep2, "CONSOLIDADO TOTAL",
                f"  VR bienes comunes  : {caso.get('vr_comun',0):>12,.2f} UF",
                f"  VR unidades        : {caso.get('vr_units',0):>12,.2f} UF",
                f"  VR TOTAL           : {vr_t:>12,.2f} UF",
                (f"  Monto asegurado    : {m_t:>12,.2f} UF" if m_t > 0 else
                 f"  Monto asegurado    :    No indicado"),
                f"  Infraseguro global : {'SÍ ⚠' if i_t else 'NO ✓'}"]
        if m_t > 0 and d_t > 0:
            lns.append(f"  Indemnización ({caso.get('danio_pct',0):.0f}%): {ind_t:>12,.2f} UF")
        if i_t and d_t > 0:
            lns.append(f"  Pérdida            : {d_t-ind_t:>12,.2f} UF")
    lns += ["", sep2,
            "Nota: Informe referencial. Verificar con tasador habilitado y póliza vigente.",
            "VUB: ingresado por el usuario. Tabla MINVU: minvu.gob.cl",
            "Normativa: DFL 251 · DS 1055 · CCom 553 · Ley 21.442 · NCG 556 CMF"]
    return "\n".join(lns)


def generar_plantilla_excel(tipo, sis, niv, vub, fg, fn, fa, p_dis, p_gg, p_ut, p_imp):
    """
    Genera un archivo Excel con la plantilla de ingreso masivo de unidades privadas.
    Los parámetros fijos (tipo, vub, factores) se precargan en cada fila como referencia.
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import io

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Unidades Privadas"

    # Estilos
    hdr_fill = PatternFill("solid", fgColor="1F3864")
    ref_fill = PatternFill("solid", fgColor="EEF3FB")
    warn_fill = PatternFill("solid", fgColor="FFF3CD")
    hdr_font = Font(name="Arial", bold=True, color="FFFFFF", size=10)
    body_font = Font(name="Arial", size=10)
    ref_font  = Font(name="Arial", size=9, italic=True, color="555555")
    center    = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left      = Alignment(horizontal="left",   vertical="center", wrap_text=True)
    thin      = Side(style="thin", color="CCCCCC")
    border    = Border(left=thin, right=thin, top=thin, bottom=thin)

    # Fila 1: Parámetros fijos (referencia)
    ws.merge_cells("A1:K1")
    ws["A1"] = (f"Parámetros fijos aplicados a todas las unidades: "
                f"Tipo={tipo}/{sis}/{niv} · VUB={vub} UF/m² · "
                f"fg={fg:.2f} · fn={fn:.2f} · fa={fa:.2f} · "
                f"Indirectos: Diseño {p_dis*100:.1f}% + GG {p_gg*100:.1f}% + "
                f"Utilidad {p_ut*100:.1f}% + Imprevistos {p_imp*100:.1f}%")
    ws["A1"].font = Font(name="Arial", size=9, italic=True, color="1F3864", bold=True)
    ws["A1"].fill = PatternFill("solid", fgColor="DBEAFE")
    ws["A1"].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    ws.row_dimensions[1].height = 32

    # Fila 2: Encabezados
    headers = [
        ("Identificación\n(obligatorio)", 22, "A"),
        ("Superficie\nm² (obligatorio)", 14, "B"),
        ("Monto asegurado\nUF (0 si no tiene)", 16, "C"),
        ("Póliza propia\n(SI/NO)", 12, "D"),
        ("Tipo\n(ref. fijo)", 12, "E"),
        ("Sistema\n(ref. fijo)", 12, "F"),
        ("Nivel\n(ref. fijo)", 12, "G"),
        ("VUB UF/m²\n(ref. fijo)", 12, "H"),
        ("fg\n(ref. fijo)", 8,  "I"),
        ("fn\n(ref. fijo)", 8,  "J"),
        ("fa\n(ref. fijo)", 8,  "K"),
    ]
    for j, (txt, w, col) in enumerate(headers, 1):
        cell = ws.cell(row=2, column=j, value=txt)
        cell.font = hdr_font
        cell.fill = hdr_fill
        cell.alignment = center
        cell.border = border
        ws.column_dimensions[get_column_letter(j)].width = w
    ws.row_dimensions[2].height = 36

    # Filas 3-52: Datos (50 unidades máx)
    for row in range(3, 53):
        # Columnas obligatorias vacías
        for col in [1, 2, 3]:
            c = ws.cell(row=row, column=col)
            c.font = body_font
            c.alignment = left if col == 1 else center
            c.border = border
        # Póliza propia
        c4 = ws.cell(row=row, column=4, value="NO")
        c4.font = body_font; c4.alignment = center; c4.border = border
        # Parámetros fijos precargados (referencia, sombreados)
        fijos = [tipo, sis, niv, vub, fg, fn, fa]
        for j, val in enumerate(fijos, 5):
            c = ws.cell(row=row, column=j, value=val)
            c.font = ref_font; c.fill = ref_fill
            c.alignment = center; c.border = border

    # Fila de instrucciones al final
    ws.merge_cells("A53:K53")
    ws["A53"] = ("INSTRUCCIONES: Complete Identificación, Superficie (m²) y Monto asegurado (UF). "
                 "Use punto para separar miles en superficies (ej: 1.250). "
                 "Las columnas sombreadas son de referencia (parámetros fijos). "
                 "No elimine filas de encabezado.")
    ws["A53"].font = Font(name="Arial", size=8, italic=True, color="666666")
    ws["A53"].fill = warn_fill
    ws["A53"].alignment = Alignment(horizontal="left", wrap_text=True)
    ws.row_dimensions[53].height = 28

    # Congelar encabezados
    ws.freeze_panes = "A3"

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def leer_excel_unidades(uploaded_file, zona, pisos, anio, aplica_iva,
                         tipo_fijo, sis_fijo, niv_fijo, vub_fijo,
                         fg_fijo, fn_fijo, fa_fijo,
                         p_dis_fijo, p_gg_fijo, p_ut_fijo, p_imp_fijo):
    """
    Lee el Excel de unidades privadas y retorna lista de dicts compatibles
    con datos_uni, usando los parámetros fijos para VUB y factores.
    """
    import openpyxl
    wb = openpyxl.load_workbook(uploaded_file, data_only=True)
    ws = wb.active
    unidades = []
    errores  = []

    for row in range(3, 53):  # filas de datos
        ident = ws.cell(row=row, column=1).value
        if not ident or str(ident).strip() == "":
            continue  # fila vacía

        sup_raw   = ws.cell(row=row, column=2).value
        monto_raw = ws.cell(row=row, column=3).value
        poliza    = str(ws.cell(row=row, column=4).value or "NO").strip().upper()

        # Parsear superficie
        sup_val = parsear_numero(str(sup_raw)) if sup_raw is not None else None
        if sup_val is None or sup_val <= 0:
            errores.append(f"Fila {row}: '{ident}' — superficie inválida ({sup_raw!r})")
            continue

        # Parsear monto
        monto_val = parsear_numero(str(monto_raw)) if monto_raw is not None else 0
        if monto_val is None:
            monto_val = 0

        unidades.append({
            "nombre":       str(ident).strip(),
            "tipo":         tipo_fijo,
            "sis":          sis_fijo,
            "niv":          niv_fijo,
            "vub":          vub_fijo,
            "sup":          int(sup_val),
            "monto":        monto_val,
            "poliza_propia": poliza == "SI",
            "zona":         zona,
            "pisos":        pisos,
            "anio":         anio,
            "aplica_iva":   aplica_iva,
            "fg_edit":      fg_fijo,
            "fn_edit":      fn_fijo,
            "fa_edit":      fa_fijo,
            "p_dis":        p_dis_fijo,
            "p_gg":         p_gg_fijo,
            "p_ut":         p_ut_fijo,
            "p_imp":        p_imp_fijo,
        })

    return unidades, errores


# ─────────────────────────────────────────────────────────
# CONFIGURACIÓN PÁGINA
# ─────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Seguro de Reconstrucción — Chile",
    page_icon="🏢", layout="centered",
    initial_sidebar_state="collapsed",
)
st.markdown("""
<style>
.stApp { max-width: 820px; margin: auto; }
.block-container { padding-top: 2rem; padding-bottom: 3rem; }
div[data-testid="stMetricValue"] { font-size: 1.3rem; }
h1 { font-size: 1.6rem !important; }
h2 { font-size: 1.2rem !important; }
</style>
""", unsafe_allow_html=True)

st.title("🏢 Calculadora de Valor de Reconstrucción")
st.caption("Seguros de inmuebles en Chile · DFL 251 · DS 1055 · CCom art. 553 · Ley 21.442 · NCG 556 CMF")

tab_calc, tab_casos, tab_como = st.tabs(["📐 Calcular", "📋 Mis casos", "ℹ️ Marco normativo"])

# ══════════════════════════════════════════════════════════
# PESTAÑA: CALCULAR
# ══════════════════════════════════════════════════════════
with tab_calc:

    # ── Identificación ──
    st.subheader("Identificación de la propiedad")
    col_n, col_d = st.columns(2)
    with col_n:
        nombre    = st.text_input("Nombre o referencia", placeholder="Ej: Edificio Torres del Parque")
    with col_d:
        direccion = st.text_input("Dirección completa", placeholder="Calle, número, comuna, región")

    # Herramienta de dirección (siempre visible si hay texto)
    widget_herramienta_direccion(direccion)

    # ── Datos generales ──
    st.subheader("Datos generales del inmueble")
    g1, g2, g3, g4 = st.columns(4)
    with g1:
        zona = st.selectbox("Zona geográfica", [""] + list(FACTOR_GEOGRAFICO.keys()),
                            format_func=lambda x: "Seleccionar..." if x == "" else x)
    with g2:
        pisos = st.number_input("N° de pisos", min_value=1, max_value=100,
                                value=None, placeholder="Ej: 12")
    with g3:
        anio = st.number_input("Año construcción", min_value=1900, max_value=2025,
                               value=None, placeholder="Ej: 2005")
    with g4:
        aplica_iva = st.checkbox("IVA (19%)", value=True)

    # ── Simulación de siniestro ──
    st.markdown("**Simulación de siniestro**")
    col_sim1, col_sim2 = st.columns(2)
    with col_sim1:
        danio_pct = st.slider(
            "% de daño a simular (0 = sin simulación)",
            min_value=0, max_value=100, value=0,
            help="Seleccione 0 para no mostrar simulación. Cualquier valor > 0 activa el cálculo.",
        )
    with col_sim2:
        perdida_real_val, perdida_real_txt, _ = input_numero(
            "O ingrese la pérdida real determinada (UF)",
            key="perdida_real",
            placeholder="Ej: 50.000",
            ayuda=(
                "Si ya determinó el monto de la pérdida (ej: peritaje), ingréselo aquí. "
                "Este campo es opcional y complementa la simulación porcentual."
            ),
            sufijo="UF",
        )
    if danio_pct == 0 and not perdida_real_val:
        st.caption("ℹ️ Sin simulación de siniestro. Seleccione un % de daño o ingrese una pérdida real para activarla.")

    datos_ok = bool(zona and pisos and anio)
    if zona and datos_ok:
        widget_tabla_refs_vub(zona)

    # ── Modo de cálculo ──
    st.subheader("¿Qué desea calcular?")
    modo = st.radio("Seleccione el alcance:", [
        "🏠  Inmueble completo — casas, locales o edificio en bloque",
        "🏛️  Solo bienes y espacios comunes — póliza de la comunidad (Ley 21.442 art. 43)",
        "🏢  Comunidad completa — bienes comunes + unidades privadas (NCG 556 CMF)",
    ])
    modo_key = ("simple" if "completo" in modo
                else "comunes" if "Solo bienes" in modo
                else "comunidad")
    st.divider()

    # ════════════════════════
    # MODO 1: INMUEBLE COMPLETO
    # ════════════════════════
    if modo_key == "simple":
        st.markdown("#### Datos del inmueble")
        if not datos_ok:
            st.info("Complete primero zona, pisos y año de construcción.")
            d = None
        else:
            d = widget_formulario_componente("s", zona, pisos, anio, aplica_iva,
                                             default_tipo="Edificio")

        if st.button("Calcular", type="primary", use_container_width=True, key="btn_s"):
            errs = [] if datos_ok else ["Complete zona, pisos y año."]
            if d: errs += validar_comp(d, "el inmueble")
            for e in errs: st.error(f"⚠️ {e}")
            if not errs:
                res = calcular_vr(d["vub"], d["sup"], zona, pisos, anio, aplica_iva, fg_override=d.get("fg_edit"), fn_override=d.get("fn_edit"), fa_override=d.get("fa_edit"), pct_diseno=d.get("p_dis"), pct_gg=d.get("p_gg"), pct_utilidad=d.get("p_ut"), pct_imprevistos=d.get("p_imp"))
                widget_resultado("Inmueble completo", res,
                                 {**d, "zona":zona,"pisos":pisos,"anio":anio}, danio_pct, perdida_real=perdida_real_val)
                caso = dict(nombre=nombre or "Sin nombre", direccion=direccion or "—",
                            zona=zona, pisos=pisos, anio=anio, danio_pct=danio_pct,
                            modo="simple",
                            comp={**d,"res":res,"zona":zona,"pisos":pisos,"anio":anio},
                            total_vr=res["vr"], total_monto=d["monto"] or 0)
                b1, b2, b3, b4 = st.columns(4)
                with b1:
                    if st.button("💾 Guardar", use_container_width=True, key="g_s"):
                        st.session_state.setdefault("casos", []).append(caso)
                        st.success("Caso guardado en 'Mis casos'.")
                with b2:
                    st.download_button("📄 Informe TXT",
                                       data=generar_informe(caso).encode(),
                                       file_name=f"informe_{(nombre or 'inmueble').replace(' ','_').lower()}.txt",
                                       mime="text/plain", use_container_width=True, key="dl_s")
                with b3:
                    docx_bytes = generar_word(caso, perdida_real=perdida_real_val)
                    if docx_bytes:
                        st.download_button("📝 Informe Word",
                                           data=docx_bytes,
                                           file_name=f"informe_{(nombre or 'inmueble').replace(' ','_').lower()}.docx",
                                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                           use_container_width=True, key="dl_s_word")

    # ════════════════════════
    # MODO 2: SOLO BIENES COMUNES
    # ════════════════════════
    elif modo_key == "comunes":
        st.markdown("#### Bienes y espacios comunes")
        st.info("**Ley 21.442, art. 43** — Seguro obligatorio de la comunidad, "
                "independiente del seguro de cada unidad privada.")
        st.caption("Incluye: estructura, fachadas, instalaciones centrales, ascensores, "
                   "subterráneos, piscina, áreas verdes, pasillos y estacionamientos comunes.")

        # ── Superficie total + distribución ──
        sup_total_bc_val, _, _ = input_numero(
            "Superficie total del edificio (m²)",
            key="sup_total_bc",
            placeholder="Ej: 6.000",
            ayuda="Ingrese la superficie total con punto como separador de miles.",
            sufijo="m²", es_entero=True,
        )
        sup_total_bc = int(sup_total_bc_val) if sup_total_bc_val else None

        # Panel de distribución — key separada para evitar conflicto con modo comunidad
        st.markdown("#### Distribución de superficies: bienes comunes vs unidades privadas")
        st.markdown(
            "No existe un porcentaje único fijado por ley — cada edificio lo define en su "
            "**Reglamento de Copropiedad**. Use la tabla siguiente como referencia:"
        )
        tabla_html_bc = """
<style>
.dist-table{width:100%;border-collapse:collapse;font-size:13px;margin-bottom:12px}
.dist-table th{background:#f0f2f6;padding:7px 10px;text-align:left;border-bottom:2px solid #ddd;font-weight:600}
.dist-table td{padding:6px 10px;border-bottom:1px solid #eee;vertical-align:top}
.dist-table tr:last-child td{background:#e8f4e8;font-weight:600}
.pct-comun{color:#1f77b4;font-weight:700;text-align:center}
.pct-units{color:#e67e22;font-weight:700;text-align:center}
.badge{display:inline-block;padding:2px 8px;border-radius:10px;font-size:11px;font-weight:600}
.b-blue{background:#dbeafe;color:#1d4ed8}
.b-orange{background:#fef3c7;color:#92400e}
.b-green{background:#d1fae5;color:#065f46}
</style>
<table class="dist-table">
<tr><th>Tipo de edificio</th><th>Bienes comunes</th><th>Unidades privadas</th><th>Descripción</th></tr>
<tr><td>🏢 Básico<br><small>2–5 pisos · sin amenidades</small></td>
    <td class="pct-comun"><span class="badge b-blue">25 – 35 %</span></td>
    <td class="pct-units"><span class="badge b-orange">65 – 75 %</span></td>
    <td><small>Pasillos, escaleras y conserjería básica. Sin subterráneos.</small></td></tr>
<tr><td>🏢 Estándar<br><small>6–15 pisos · 1 subterráneo</small></td>
    <td class="pct-comun"><span class="badge b-blue">35 – 45 %</span></td>
    <td class="pct-units"><span class="badge b-orange">55 – 65 %</span></td>
    <td><small>Piscina, gimnasio, sala multiuso y 1 nivel de estacionamientos.</small></td></tr>
<tr><td>🏢 Alto estándar<br><small>15–25 pisos · 2 subterráneos</small></td>
    <td class="pct-comun"><span class="badge b-blue">45 – 55 %</span></td>
    <td class="pct-units"><span class="badge b-orange">45 – 55 %</span></td>
    <td><small>Múltiples amenidades, lobby amplio, 2 niveles de estacionamientos.</small></td></tr>
<tr><td>🏢 Premium / Torre<br><small>25+ pisos · 3+ subterráneos</small></td>
    <td class="pct-comun"><span class="badge b-blue">55 – 70 %</span></td>
    <td class="pct-units"><span class="badge b-orange">30 – 45 %</span></td>
    <td><small>Todas las amenidades, lobby doble altura, spa, múltiples subterráneos.</small></td></tr>
<tr><td>📋 Referencia Ley 21.442</td>
    <td class="pct-comun"><span class="badge b-green">50–70%</span><br><small>del VR total</small></td>
    <td class="pct-units">—</td>
    <td><small>Bienes comunes = 50–70% del <strong>valor</strong> de reconstrucción.</small></td></tr>
</table>
<p style="font-size:11px;color:#888">Fuentes: Edifito / Ley 21.442 · ComunidadFeliz · OGUC art. 5.1.11 · Práctica de mercado 2025</p>
"""
        st.markdown(tabla_html_bc, unsafe_allow_html=True)

        # Slider con key ÚNICA para este modo
        pct_bc_pct = st.slider(
            "Seleccione el % de superficie de bienes comunes para este edificio",
            min_value=20, max_value=70, value=40, step=1,
            help="Use la tabla anterior como guía. Para el valor exacto consulte el Reglamento de Copropiedad.",
            key="pct_bc_slider_solo",
        )
        pct_comun_bc = pct_bc_pct / 100

        # Barra visual
        bar_bc = f"""
<div style="margin:8px 0 4px 0">
  <div style="display:flex;height:28px;border-radius:6px;overflow:hidden;border:1px solid #ddd">
    <div style="width:{pct_bc_pct}%;background:#1f77b4;display:flex;align-items:center;
                justify-content:center;color:white;font-size:12px;font-weight:600;min-width:30px">
      {pct_bc_pct}%</div>
    <div style="width:{100-pct_bc_pct}%;background:#e67e22;display:flex;align-items:center;
                justify-content:center;color:white;font-size:12px;font-weight:600;min-width:30px">
      {100-pct_bc_pct}%</div>
  </div>
  <div style="display:flex;gap:20px;margin-top:5px;font-size:12px">
    <span style="color:#1f77b4">■ Bienes comunes: <strong>{pct_bc_pct}%</strong></span>
    <span style="color:#e67e22">■ Unidades privadas: <strong>{100-pct_bc_pct}%</strong></span>
  </div>
</div>
"""
        st.markdown(bar_bc, unsafe_allow_html=True)

        # Calcular superficies
        sup_comun_bc = round(sup_total_bc * pct_comun_bc) if sup_total_bc else None
        sup_units_bc = (sup_total_bc - sup_comun_bc) if (sup_total_bc and sup_comun_bc) else None

        # Desglose en m² (si hay superficie total)
        if sup_total_bc and sup_comun_bc:
            desglose_bc = f"""
<div style="display:grid;grid-template-columns:1fr 1fr 1fr;gap:10px;margin:10px 0">
  <div style="background:#f8f9fa;border-radius:8px;padding:12px;border-left:4px solid #888;text-align:center">
    <div style="font-size:11px;color:#666;margin-bottom:4px">Superficie total</div>
    <div style="font-size:20px;font-weight:600">{fmt_miles(sup_total_bc, 0)} m²</div>
  </div>
  <div style="background:#dbeafe;border-radius:8px;padding:12px;border-left:4px solid #1f77b4;text-align:center">
    <div style="font-size:11px;color:#1d4ed8;margin-bottom:4px">Bienes comunes ({pct_bc_pct}%)</div>
    <div style="font-size:20px;font-weight:600;color:#1d4ed8">{fmt_miles(sup_comun_bc, 0)} m²</div>
  </div>
  <div style="background:#fef3c7;border-radius:8px;padding:12px;border-left:4px solid #e67e22;text-align:center">
    <div style="font-size:11px;color:#92400e;margin-bottom:4px">Unidades privadas ({100-pct_bc_pct}%)</div>
    <div style="font-size:20px;font-weight:600;color:#92400e">{fmt_miles(sup_units_bc, 0)} m²</div>
  </div>
</div>
"""
            st.markdown(desglose_bc, unsafe_allow_html=True)
            st.success(
                f"✅ El cálculo usará **{fmt_m2(sup_comun_bc)}** × VUB "
                f"({pct_bc_pct}% de {fmt_m2(sup_total_bc)} totales) como superficie de bienes comunes."
            )
        else:
            st.caption("⚠️ Ingrese la superficie total del edificio para ver el desglose en m².")

        if not datos_ok:
            st.info("Complete primero zona, pisos y año.")
            d = None
        else:
            st.markdown("---")
            st.markdown("##### Datos del componente — Bienes y espacios comunes")

            col_tipo1, col_tipo2, col_tipo3 = st.columns(3)
            with col_tipo1:
                tipo_bc = st.selectbox("Tipo", list(SISTEMAS_POR_TIPO.keys()),
                                       index=list(SISTEMAS_POR_TIPO.keys()).index("Comunidad"),
                                       key="bc_tipo")
            with col_tipo2:
                sis_bc = st.selectbox("Sistema constructivo",
                                      SISTEMAS_POR_TIPO[tipo_bc], key="bc_sis")
            with col_tipo3:
                niv_bc = st.selectbox("Nivel de terminaciones",
                                      NIVELES_POR_TS[(tipo_bc, sis_bc)], key="bc_niv",
                                      help="Básico = sin lujos · Medio = estándar · Alto = premium")

            # VUB con tabla de referencia
            vub_bc = widget_vub("bc", zona, tipo_bc, sis_bc, niv_bc)
            st.caption("ℹ️ **El VUB se ingresa sin IVA.** El IVA (19%) se aplica al final sobre el subtotal.")

            # Factores — siempre visibles, no dentro de expander
            st.markdown("**Factores del cálculo**")
            st.caption("Calculados automáticamente según zona, año y pisos. Puede ajustarlos si dispone de datos precisos.")
            ff1, ff2, ff3 = st.columns(3)
            with ff1:
                bc_fg = st.number_input("Factor geográfico",
                                        min_value=0.50, max_value=2.00,
                                        value=round(FACTOR_GEOGRAFICO.get(zona, 1.0), 2),
                                        step=0.01, format="%.2f", key="bc_fg",
                                        help="Metropolitana 1,05 · Intermedia 1,00 · Aislada 1,15")
            with ff2:
                bc_fn = st.number_input(f"Factor normativo (año {anio})",
                                        min_value=0.50, max_value=2.00,
                                        value=round(factor_normativo(anio), 2),
                                        step=0.01, format="%.2f", key="bc_fn",
                                        help="<1985: 1,15 · 1985-2000: 1,10 · 2001-2010: 1,05 · >2010: 1,00")
            with ff3:
                bc_fa = st.number_input(f"Factor altura ({pisos} pisos)",
                                        min_value=0.50, max_value=2.00,
                                        value=round(factor_altura(pisos), 2),
                                        step=0.01, format="%.2f", key="bc_fa",
                                        help="1-2p: 1,00 · 3-5p: 1,05 · 6-10p: 1,10 · 11+: 1,15")

            st.markdown("**Costos indirectos (% sobre costo directo)**")
            fp1, fp2, fp3, fp4 = st.columns(4)
            with fp1:
                bc_pdis = st.number_input("Diseño (%)", min_value=0.0, max_value=30.0,
                                          value=3.0, step=0.5, format="%.1f", key="bc_pdis") / 100
            with fp2:
                bc_pgg = st.number_input("Gastos grales. (%)", min_value=0.0, max_value=30.0,
                                         value=6.0, step=0.5, format="%.1f", key="bc_pgg") / 100
            with fp3:
                bc_put = st.number_input("Utilidad contr. (%)", min_value=0.0, max_value=30.0,
                                         value=12.0, step=0.5, format="%.1f", key="bc_put") / 100
            with fp4:
                bc_pimp = st.number_input("Imprevistos (%)", min_value=0.0, max_value=30.0,
                                          value=10.0, step=0.5, format="%.1f", key="bc_pimp") / 100
            total_ind_bc = (bc_pdis + bc_pgg + bc_put + bc_pimp) * 100
            st.caption(f"Total costos indirectos: **{total_ind_bc:.1f}%** sobre el costo directo.")

            # Preview del cálculo antes de calcular
            if vub_bc and sup_comun_bc:
                cd_prev = sup_comun_bc * vub_bc * bc_fg * bc_fn * bc_fa
                ci_prev = cd_prev * (bc_pdis + bc_pgg + bc_put + bc_pimp)
                st_prev = cd_prev + ci_prev
                iv_prev = st_prev * TASA_IVA if aplica_iva else 0
                vr_prev = st_prev + iv_prev
                st.info(
                    f"**Vista previa del cálculo:** "
                    f"{fmt_m2(sup_comun_bc)} × {vub_bc:.1f} UF/m² × {bc_fg:.2f} × {bc_fn:.2f} × {bc_fa:.2f} "
                    f"= costo directo **{fmt_uf(cd_prev)}** → "
                    f"+ ind. {total_ind_bc:.0f}% + IVA 19% = **{fmt_uf(vr_prev)}**"
                )

            monto_bc_val, _, _ = input_numero(
                "Monto asegurado en póliza (UF)", key="bc_monto",
                placeholder="Ej: 227.500",
                ayuda="Monto de la póliza de la comunidad para bienes comunes. Ingrese 0 si no hay seguro.",
                sufijo="UF",
            )
            monto_bc = monto_bc_val if monto_bc_val is not None else 0
            d = {"tipo": tipo_bc, "sis": sis_bc, "niv": niv_bc,
                 "vub": vub_bc, "sup": sup_comun_bc, "monto": monto_bc,
                 "sup_total": sup_total_bc, "pct_comun": pct_comun_bc,
                 "zona": zona, "pisos": pisos, "anio": anio, "aplica_iva": aplica_iva,
                 "fg_edit": bc_fg, "fn_edit": bc_fn, "fa_edit": bc_fa,
                 "p_dis": bc_pdis, "p_gg": bc_pgg, "p_ut": bc_put, "p_imp": bc_pimp}

        if st.button("Calcular bienes comunes", type="primary",
                     use_container_width=True, key="btn_bc"):
            errs = [] if datos_ok else ["Complete zona, pisos y año."]
            if d:
                if not d.get("vub"):   errs.append("Ingrese el VUB (UF/m²).")
                if not d.get("sup"):   errs.append("Ingrese la superficie total del edificio y seleccione el porcentaje de bienes comunes.")
                if d.get("monto") is None: errs.append("Ingrese el monto asegurado (puede ser 0).")
            for e in errs: st.error(f"⚠️ {e}")
            if not errs:
                res = calcular_vr(d["vub"], d["sup"], zona, pisos, anio, aplica_iva, fg_override=d.get("fg_edit"), fn_override=d.get("fn_edit"), fa_override=d.get("fa_edit"), pct_diseno=d.get("p_dis"), pct_gg=d.get("p_gg"), pct_utilidad=d.get("p_ut"), pct_imprevistos=d.get("p_imp"))

                # Nota explicativa del desglose usado
                nota_bc = (
                    f"Asegurado: la comunidad (Ley 21.442 art. 43 — OBLIGATORIO) · "
                    f"Superficie bienes comunes: {fmt_m2(d['sup'])} "
                    f"({pct_bc_pct}% de {fmt_m2(sup_total_bc)} totales)"
                )
                widget_resultado("Bienes y espacios comunes", res,
                                 {**d, "zona":zona, "pisos":pisos, "anio":anio},
                                 danio_pct, perdida_real=perdida_real_val, nota=nota_bc)
                caso = dict(nombre=nombre or "Sin nombre", direccion=direccion or "—",
                            zona=zona, pisos=pisos, anio=anio, danio_pct=danio_pct,
                            modo="comunes",
                            comp={**d, "res":res, "zona":zona, "pisos":pisos, "anio":anio},
                            total_vr=res["vr"], total_monto=d["monto"] or 0)
                b1, b2, b3 = st.columns(3)
                with b1:
                    if st.button("💾 Guardar", use_container_width=True, key="g_bc"):
                        st.session_state.setdefault("casos", []).append(caso)
                        st.success("Guardado.")
                with b2:
                    st.download_button("📄 Informe TXT",
                                       data=generar_informe(caso).encode(),
                                       file_name=f"informe_{(nombre or 'comunes').replace(' ','_').lower()}.txt",
                                       mime="text/plain", use_container_width=True, key="dl_bc")
                with b3:
                    docx_bytes_bc = generar_word(caso, perdida_real=perdida_real_val)
                    if docx_bytes_bc:
                        st.download_button("📝 Informe Word",
                                           data=docx_bytes_bc,
                                           file_name=f"informe_{(nombre or 'comunes').replace(' ','_').lower()}.docx",
                                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                           use_container_width=True, key="dl_bc_word")

    # ════════════════════════
    # MODO 3: COMUNIDAD COMPLETA
    # ════════════════════════
    else:
        st.info(
            "**NCG 556 CMF (dic. 2025)** — La póliza colectiva se estructura en dos bloques separados: "
            "**Bloque 1** bienes comunes (asegurado: la comunidad) · "
            "**Bloque 2** unidades privadas (asegurado: cada copropietario)."
        )

        # ── Superficie total ──
        st.markdown("---")
        sup_total_val, _, _ = input_numero(
            "Superficie total del edificio (m²)",
            key="sup_total_edificio",
            placeholder="Ej: 12.000",
            ayuda="Suma de TODOS los pisos + subterráneos. Use punto como separador de miles.",
            sufijo="m²", es_entero=True,
        )
        sup_total = int(sup_total_val) if sup_total_val else None

        # ── Panel de distribución (debajo de Superficie total) ──
        pct_comun, sup_comun_calc, sup_units_calc = widget_distribucion_superficies(sup_total)

        # ── Bloque 1: Bienes comunes ──
        st.markdown("---")
        st.markdown("#### Bloque 1 — Bienes y espacios comunes")
        st.caption("Estructura, fachadas, instalaciones centrales, ascensores, "
                   "subterráneos y toda área de dominio común.")
        if sup_total and sup_comun_calc:
            st.success(
                f"✅ Se usarán **{fmt_m2(sup_comun_calc)}** para bienes comunes "
                f"({int(pct_comun*100)}% de {fmt_m2(sup_total)} totales). "
                f"Puede ajustar este valor en el campo Superficie si lo conoce exactamente."
            )
        if datos_ok:
            d_comun = widget_formulario_componente(
                "c_bc", zona, pisos, anio, aplica_iva,
                default_tipo="Comunidad", label_tipo="Tipo (bienes comunes)",
                sup_sugerida=sup_comun_calc,
                mostrar_dist=True,
                pct_comun=pct_comun,
                pct_pct=int(pct_comun * 100),
                sup_total=sup_total,
            )
        else:
            st.info("Complete zona, pisos y año de construcción.")
            d_comun = None

        # ── Bloque 2: Unidades privadas ──
        st.markdown("---")
        st.markdown("#### Bloque 2 — Unidades privadas")
        st.caption("VUB con tipo **Depto** — valor de la unidad habitable sin incluir áreas comunes.")

        # ── Bloque 2: Unidades privadas ──
        st.markdown("---")
        st.markdown("#### Bloque 2 — Unidades privadas")
        st.caption("VUB con tipo **Depto** — valor de la unidad habitable sin incluir áreas comunes.")

        incluir_uni = st.checkbox(
            "Incluir unidades privadas en este análisis", value=True,
            help="Desmarque si las unidades tienen seguros individuales separados.",
        )
        datos_uni = []

        if incluir_uni and datos_ok:

            # ── Panel de parámetros fijos ──
            st.markdown("##### ⚙️ Parámetros comunes para todas las unidades")
            st.caption(
                "Defina aquí los parámetros que se aplicarán a **todas** las unidades. "
                "En el ingreso individual puede sobreescribir estos valores por unidad."
            )
            pu1, pu2, pu3 = st.columns(3)
            with pu1:
                uni_tipo = st.selectbox("Tipo de unidad", list(SISTEMAS_POR_TIPO.keys()),
                                        index=list(SISTEMAS_POR_TIPO.keys()).index("Depto"),
                                        key="uni_tipo_fijo")
            with pu2:
                uni_sis = st.selectbox("Sistema constructivo",
                                       SISTEMAS_POR_TIPO[uni_tipo], key="uni_sis_fijo")
            with pu3:
                uni_niv = st.selectbox("Nivel de terminaciones",
                                       NIVELES_POR_TS[(uni_tipo, uni_sis)],
                                       key="uni_niv_fijo",
                                       help="Básico = sin lujos · Medio = estándar · Alto = premium")

            uni_vub = widget_vub("uni_fijo", zona, uni_tipo, uni_sis, uni_niv)
            st.caption("ℹ️ **VUB sin IVA.** El IVA 19% se aplica al final sobre el subtotal.")

            st.markdown("**Factores del cálculo (comunes a todas las unidades)**")
            pf1, pf2, pf3 = st.columns(3)
            with pf1:
                uni_fg = st.number_input("Factor geográfico", min_value=0.50, max_value=2.00,
                                         value=round(float(FACTOR_GEOGRAFICO.get(zona, 1.0)),2),
                                         step=0.01, format="%.2f", key="uni_fg_fijo")
            with pf2:
                uni_fn = st.number_input(f"Factor normativo (año {anio})",
                                         min_value=0.50, max_value=2.00,
                                         value=round(factor_normativo(anio),2),
                                         step=0.01, format="%.2f", key="uni_fn_fijo")
            with pf3:
                uni_fa = st.number_input(f"Factor altura ({pisos} pisos)",
                                         min_value=0.50, max_value=2.00,
                                         value=round(factor_altura(pisos),2),
                                         step=0.01, format="%.2f", key="uni_fa_fijo")

            st.markdown("**Costos indirectos**")
            pi1, pi2, pi3, pi4 = st.columns(4)
            with pi1:
                uni_pdis = st.number_input("Diseño (%)", min_value=0.0, max_value=30.0,
                                           value=3.0, step=0.5, format="%.1f", key="uni_pdis") / 100
            with pi2:
                uni_pgg = st.number_input("Gastos grales. (%)", min_value=0.0, max_value=30.0,
                                          value=6.0, step=0.5, format="%.1f", key="uni_pgg") / 100
            with pi3:
                uni_put = st.number_input("Utilidad contr. (%)", min_value=0.0, max_value=30.0,
                                          value=12.0, step=0.5, format="%.1f", key="uni_put") / 100
            with pi4:
                uni_pimp = st.number_input("Imprevistos (%)", min_value=0.0, max_value=30.0,
                                           value=10.0, step=0.5, format="%.1f", key="uni_pimp") / 100
            st.caption(f"Total indirectos: **{(uni_pdis+uni_pgg+uni_put+uni_pimp)*100:.1f}%**")

            params_fijos_ok = bool(uni_vub)

            st.markdown("---")
            st.markdown("##### Ingreso de unidades")

            # ── Sugerencia de superficie ──
            sup_uni_sug = None
            if sup_units_calc and st.session_state.get("n_uni", 1) > 0:
                sup_uni_sug = round(sup_units_calc / st.session_state.get("n_uni", 1))

            # ── Tabs: Individual / Masivo ──
            tab_ind, tab_mas = st.tabs(["📋 Ingreso individual", "📊 Ingreso masivo por Excel"])

            # ════ TAB INDIVIDUAL ════
            with tab_ind:
                ca, cr = st.columns(2)
                with ca:
                    if st.button("➕ Agregar unidad", use_container_width=True, key="add_uni"):
                        st.session_state.n_uni = st.session_state.get("n_uni", 1) + 1
                with cr:
                    if st.button("➖ Quitar última", use_container_width=True,
                                 key="del_uni",
                                 disabled=st.session_state.get("n_uni", 1) <= 1):
                        st.session_state.n_uni = st.session_state.get("n_uni", 1) - 1

                if sup_units_calc:
                    n_u = st.session_state.get("n_uni", 1)
                    sup_uni_sug = round(sup_units_calc / n_u)
                    st.info(
                        f"Superficie unidades privadas: **{fmt_m2(sup_units_calc)}** "
                        f"({100-int(pct_comun*100)}% de {fmt_m2(sup_total)}) · "
                        f"Promedio por unidad: **{fmt_m2(sup_uni_sug)}** "
                        f"({fmt_m2(sup_units_calc)} ÷ {n_u} unidades)"
                    )

                datos_uni_ind = []
                for i in range(st.session_state.get("n_uni", 1)):
                    with st.expander(f"Unidad {i+1}", expanded=(i == 0)):
                        nom_u = st.text_input("Identificación", key=f"u_{i}_nom",
                                              placeholder="Ej: Depto 501, Local 2")
                        st.caption(
                            "Los parámetros fijos (tipo, VUB, factores) se aplican desde el panel superior. "
                            "Solo ingrese superficie y monto para esta unidad."
                        )
                        sup_u_val, _, _ = input_numero(
                            "Superficie (m²)", key=f"u_{i}_sup",
                            placeholder="Ej: 75" if not sup_uni_sug else f"Ej: {fmt_miles(sup_uni_sug,0)}",
                            sufijo="m²", es_entero=True,
                        )
                        sup_u = int(sup_u_val) if sup_u_val else sup_uni_sug
                        monto_u_val, _, _ = input_numero(
                            "Monto asegurado (UF)", key=f"u_{i}_monto",
                            placeholder="Ej: 2.500", sufijo="UF",
                        )
                        monto_u = monto_u_val if monto_u_val is not None else 0
                        poliza_u = st.checkbox("Póliza propia vigente", key=f"u_{i}_prop")

                        datos_uni_ind.append({
                            "nombre": nom_u or f"Unidad {i+1}",
                            "tipo": uni_tipo, "sis": uni_sis, "niv": uni_niv,
                            "vub": uni_vub, "sup": sup_u, "monto": monto_u,
                            "poliza_propia": poliza_u,
                            "zona": zona, "pisos": pisos, "anio": anio,
                            "aplica_iva": aplica_iva,
                            "fg_edit": uni_fg, "fn_edit": uni_fn, "fa_edit": uni_fa,
                            "p_dis": uni_pdis, "p_gg": uni_pgg,
                            "p_ut": uni_put, "p_imp": uni_pimp,
                        })

                datos_uni = datos_uni_ind

            # ════ TAB MASIVO ════
            with tab_mas:
                st.markdown(
                    "**Flujo:** ① Descargue la plantilla → ② Complétela en Excel → "
                    "③ Cargue el archivo → ④ Verifique el resumen → ⑤ Calcule"
                )
                st.caption(
                    "Los parámetros fijos definidos en el panel superior "
                    "(tipo, VUB, factores) se aplican automáticamente a todas las unidades. "
                    "Solo debe completar: **Identificación**, **Superficie (m²)** y **Monto asegurado (UF)**."
                )

                if params_fijos_ok:
                    # Descarga de plantilla
                    plantilla_bytes = generar_plantilla_excel(
                        uni_tipo, uni_sis, uni_niv, uni_vub,
                        uni_fg, uni_fn, uni_fa,
                        uni_pdis, uni_pgg, uni_put, uni_pimp,
                    )
                    st.download_button(
                        "📥 Descargar plantilla Excel",
                        data=plantilla_bytes,
                        file_name="plantilla_unidades_privadas.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True, key="dl_plantilla",
                        help="Descargue, complete y luego cargue el archivo completado."
                    )
                else:
                    st.warning("⚠️ Ingrese primero el VUB en el panel de parámetros comunes para generar la plantilla.")

                st.markdown("---")
                uploaded_excel = st.file_uploader(
                    "③ Cargar Excel completado",
                    type=["xlsx"],
                    key="upload_excel_uni",
                    help="Cargue el archivo Excel con las unidades completadas.",
                )

                if uploaded_excel and params_fijos_ok:
                    unis_excel, errores_excel = leer_excel_unidades(
                        uploaded_excel, zona, pisos, anio, aplica_iva,
                        uni_tipo, uni_sis, uni_niv, uni_vub,
                        uni_fg, uni_fn, uni_fa,
                        uni_pdis, uni_pgg, uni_put, uni_pimp,
                    )
                    if errores_excel:
                        for err in errores_excel:
                            st.error(f"⚠️ {err}")
                    if unis_excel:
                        st.success(f"✅ {len(unis_excel)} unidades leídas correctamente.")
                        # Tabla de resumen
                        resumen = []
                        for u in unis_excel:
                            cd_p = u["sup"] * u["vub"] * u["fg_edit"] * u["fn_edit"] * u["fa_edit"]
                            ci_p = cd_p * (u["p_dis"]+u["p_gg"]+u["p_ut"]+u["p_imp"])
                            st_p = cd_p + ci_p
                            vr_p = st_p * (1 + TASA_IVA) if aplica_iva else st_p
                            resumen.append({
                                "Identificación": u["nombre"],
                                "Superficie": fmt_m2(u["sup"]),
                                "Monto asegurado": fmt_uf(u["monto"]) if u["monto"] > 0 else "—",
                                "VR estimado": fmt_uf(vr_p),
                                "Póliza propia": "Sí" if u["poliza_propia"] else "No",
                            })
                        st.dataframe(pd.DataFrame(resumen), use_container_width=True, hide_index=True)
                        datos_uni = unis_excel
                elif uploaded_excel and not params_fijos_ok:
                    st.warning("⚠️ Ingrese el VUB en el panel de parámetros comunes antes de cargar el Excel.")

        elif incluir_uni and not datos_ok:
            st.info("Complete los datos generales para habilitar las unidades.")

        st.divider()
        if st.button("Calcular comunidad completa", type="primary",
                     use_container_width=True, key="btn_com"):
            errs = [] if datos_ok else ["Complete zona, pisos y año."]
            if d_comun:
                errs += validar_comp(d_comun, "bienes comunes")
            else:
                errs += ["Complete los datos de bienes comunes."]
            if incluir_uni:
                for i, du in enumerate(datos_uni, 1):
                    errs += validar_comp(du, f"unidad {i}")
            for e in errs:
                st.error(f"⚠️ {e}")

            if not errs:
                res_c = calcular_vr(d_comun["vub"], d_comun["sup"], zona, pisos, anio, aplica_iva, fg_override=d_comun.get("fg_edit"), fn_override=d_comun.get("fn_edit"), fa_override=d_comun.get("fa_edit"), pct_diseno=d_comun.get("p_dis"), pct_gg=d_comun.get("p_gg"), pct_utilidad=d_comun.get("p_ut"), pct_imprevistos=d_comun.get("p_imp"))
                comp_c = {**d_comun, "res":res_c, "zona":zona, "pisos":pisos, "anio":anio}

                units_calc = []
                for du in datos_uni:
                    r_u = calcular_vr(du["vub"], du["sup"], zona, pisos, anio, aplica_iva, fg_override=du.get("fg_edit"), fn_override=du.get("fn_edit"), fa_override=du.get("fa_edit"), pct_diseno=du.get("p_dis"), pct_gg=du.get("p_gg"), pct_utilidad=du.get("p_ut"), pct_imprevistos=du.get("p_imp"))
                    units_calc.append({**du, "res":r_u, "zona":zona, "pisos":pisos, "anio":anio})

                vr_c   = res_c["vr"]
                vr_u   = sum(u["res"]["vr"] for u in units_calc)
                vr_t   = vr_c + vr_u
                m_c    = d_comun["monto"] or 0
                m_u    = sum(u.get("monto") or 0 for u in datos_uni)
                m_t    = m_c + m_u
                r_t, i_t = evaluar(m_t, vr_t)
                # Daño consolidado: pérdida real si se ingresó, sino % simulado
                if perdida_real_val and perdida_real_val > 0:
                    d_t = perdida_real_val
                    origen_d_t = f"pérdida real ingresada ({fmt_uf(perdida_real_val)})"
                elif danio_pct > 0:
                    d_t = vr_t * (danio_pct / 100)
                    origen_d_t = f"simulación {danio_pct}% del VR total"
                else:
                    d_t = None
                    origen_d_t = None
                ind_t = indemn(d_t, m_t, vr_t) if d_t is not None else None

                st.divider()
                st.subheader("Resultados")

                # Resumen consolidado
                st.markdown("##### Resumen consolidado")
                if m_t <= 0:
                    st.info("ℹ️ Sin monto asegurado. El valor calculado indica cuánto debería asegurarse.")
                elif i_t:
                    st.warning(f"⚠️ **Infraseguro global.** Cobertura: **{r_t*100:.1f}%** "
                               f"— Brecha: **{fmt_uf(vr_t-m_t)}**")
                else:
                    st.success(f"✅ Cobertura global adecuada ({r_t*100:.1f}%)")

                t1, t2, t3, t4 = st.columns(4)
                t1.metric("VR total comunidad", fmt_uf(vr_t))
                t2.metric("Bienes comunes",     fmt_uf(vr_c))
                t3.metric("Unidades privadas",  fmt_uf(vr_u))
                t4.metric("Cobertura global",
                          f"{r_t*100:.1f}%" if m_t > 0 else "—",
                          delta=f"{(r_t-1)*100:.1f}%" if m_t > 0 else None,
                          delta_color="normal" if not i_t else "inverse")
                if m_t > 0:
                    st.progress(min(r_t, 1.0), text=f"Cobertura global: {r_t*100:.1f}%")

                if d_t is not None and d_t > 0:
                    with st.expander(f"🔥 Simulación total — {origen_d_t}"):
                        sc1, sc2, sc3 = st.columns(3)
                        sc1.metric("Daño / Pérdida total", fmt_uf(d_t))
                        sc2.metric("Indemnización global",
                                   fmt_uf(ind_t) if m_t > 0 else "—")
                        if i_t and m_t > 0:
                            sc3.metric("Pérdida no cubierta",
                                       fmt_uf(d_t - ind_t), delta_color="inverse")
                            st.warning(
                                f"**Art. 553 CCom:** la comunidad recibiría **{fmt_uf(ind_t)}** "
                                f"en vez de **{fmt_uf(d_t)}**. "
                                f"Pérdida no cubierta: **{fmt_uf(d_t - ind_t)}**."
                            )

                # Tabla comparativa
                st.markdown("---")
                st.markdown("##### Tabla comparativa por componente")
                filas = []
                r_c2, i_c2 = evaluar(m_c, vr_c)
                filas.append({
                    "Componente":"Bienes comunes","Asegurado":"Comunidad",
                    "Sup. m²":fmt_miles(d_comun['sup'] or 0, 0),"VUB":f"{d_comun['vub']:.1f}",
                    "VR (UF)":fmt_miles(vr_c, 2),
                    "Asegurado (UF)":fmt_miles(m_c, 2) if m_c > 0 else "—",
                    "Cobertura":f"{r_c2*100:.1f}%" if m_c > 0 else "—",
                    "Estado":"⚠️" if i_c2 else ("✅" if m_c > 0 else "ℹ️"),
                })
                for u in units_calc:
                    vr_u2 = u["res"]["vr"]
                    m_u2  = u.get("monto") or 0
                    r_u2, i_u2 = evaluar(m_u2, vr_u2)
                    pp = " (póliza propia)" if u.get("poliza_propia") else ""
                    filas.append({
                        "Componente": u.get("nombre") or "Unidad",
                        "Asegurado": f"Copropietario{pp}",
                        "Sup. m²": fmt_miles(u['sup'] or 0, 0), "VUB": f"{u['vub']:.1f}",
                        "VR (UF)": fmt_miles(vr_u2, 2),
                        "Asegurado (UF)": fmt_miles(m_u2, 2) if m_u2 > 0 else "—",
                        "Cobertura": f"{r_u2*100:.1f}%" if m_u2 > 0 else "—",
                        "Estado": "⚠️" if i_u2 else ("✅" if m_u2 > 0 else "ℹ️"),
                    })
                filas.append({
                    "Componente":"TOTAL","Asegurado":"—",
                    "Sup. m²": fmt_miles((d_comun['sup'] or 0)+sum(u.get('sup') or 0 for u in datos_uni), 0),
                    "VUB":"—","VR (UF)":fmt_miles(vr_t, 2),
                    "Asegurado (UF)":fmt_miles(m_t, 2) if m_t > 0 else "—",
                    "Cobertura":f"{r_t*100:.1f}%" if m_t > 0 else "—",
                    "Estado":"⚠️" if i_t else ("✅" if m_t > 0 else "ℹ️"),
                })
                st.dataframe(pd.DataFrame(filas), use_container_width=True, hide_index=True)

                # Detalle por componente
                st.markdown("---")
                st.markdown("##### Detalle por componente")
                widget_resultado("Bienes y espacios comunes", res_c, comp_c, danio_pct,
                                 nota="Asegurado: la comunidad (Ley 21.442 art. 43 — OBLIGATORIO)",
                                 expanded=True)
                for i, u in enumerate(units_calc, 1):
                    lbl = u.get("nombre") or f"Unidad {i}"
                    nota_u = ("Póliza propia — puede renunciar a cobertura colectiva"
                              if u.get("poliza_propia")
                              else "Asegurado: copropietario (NCG 556 Bloque 2)")
                    widget_resultado(lbl, u["res"], u, danio_pct,
                                     perdida_real=perdida_real_val,
                                     nota=nota_u, expanded=(i == 1))

                # Guardar / Exportar
                pct_pct_val = int(pct_comun * 100)
                caso = dict(
                    nombre=nombre or "Sin nombre", direccion=direccion or "—",
                    zona=zona, pisos=pisos, anio=anio, danio_pct=danio_pct,
                    modo="comunidad",
                    comp_comun=comp_c,
                    unidades=[{**u, "nombre": u.get("nombre") or f"Unidad {j+1}"}
                               for j, u in enumerate(units_calc)],
                    vr_comun=vr_c, vr_units=vr_u, total_vr=vr_t, total_monto=m_t,
                    desglose={
                        "sup_total": sup_total, "pct_pct": pct_pct_val,
                        "sup_comun": sup_comun_calc, "sup_units": sup_units_calc,
                    } if sup_total else None,
                )
                b1, b2, b3 = st.columns(3)
                with b1:
                    if st.button("💾 Guardar", use_container_width=True, key="g_com"):
                        st.session_state.setdefault("casos", []).append(caso)
                        st.success("Guardado en 'Mis casos'.")
                with b2:
                    st.download_button(
                        "📄 Informe TXT",
                        data=generar_informe(caso).encode(),
                        file_name=f"informe_{(nombre or 'comunidad').replace(' ','_').lower()}.txt",
                        mime="text/plain", use_container_width=True, key="dl_com",
                    )
                with b3:
                    docx_bytes_com = generar_word(caso, perdida_real=perdida_real_val)
                    if docx_bytes_com:
                        st.download_button(
                            "📝 Informe Word",
                            data=docx_bytes_com,
                            file_name=f"informe_{(nombre or 'comunidad').replace(' ','_').lower()}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True, key="dl_com_word",
                        )

# ══════════════════════════════════════════════════════════
# PESTAÑA: MIS CASOS
# ══════════════════════════════════════════════════════════
with tab_casos:
    casos = st.session_state.get("casos", [])
    if not casos:
        st.info("Aún no tiene casos guardados.")
    else:
        st.caption(f"{len(casos)} caso{'s' if len(casos) > 1 else ''} guardado{'s' if len(casos) > 1 else ''}")
        modos_lbl = {"simple":"Completo","comunes":"Bienes comunes","comunidad":"Comunidad"}
        for i, c in enumerate(casos):
            vr_c = c["total_vr"]
            m_c  = c["total_monto"]
            r_c, inf_c = evaluar(m_c, vr_c)
            estado = "⚠️ Infraseguro" if inf_c else ("✅ Cubierto" if m_c > 0 else "ℹ️ Sin seguro")
            with st.expander(
                f"{estado}  |  {c['nombre']}  —  {fmt_uf(vr_c)}  [{modos_lbl.get(c['modo'],'—')}]"
            ):
                cc1, cc2, cc3 = st.columns(3)
                cc1.metric("Valor de reconstrucción", fmt_uf(vr_c))
                cc2.metric("Monto asegurado", fmt_uf(m_c) if m_c > 0 else "No indicado")
                cc3.metric("Cobertura", f"{r_c*100:.1f}%" if m_c > 0 else "—")
                st.caption(f"{c.get('zona','—')} · {c.get('pisos','—')} pisos · año {c.get('anio','—')}")
                st.download_button(
                    "📄 Descargar informe",
                    data=generar_informe(c).encode(),
                    file_name=f"informe_{c['nombre'].replace(' ','_').lower()}.txt",
                    mime="text/plain", key=f"dl_caso_{i}",
                )
        if st.button("🗑️ Limpiar todos los casos"):
            st.session_state.casos = []
            st.rerun()

# ══════════════════════════════════════════════════════════
# PESTAÑA: MARCO NORMATIVO
# ══════════════════════════════════════════════════════════
with tab_como:
    st.subheader("¿Qué es el valor de reconstrucción?")
    st.markdown("""
Es el costo real de volver a construir la propiedad **desde cero**.
Es el monto que debe quedar cubierto en la póliza. Si la póliza cubre menos, hay **infraseguro**
y la compañía pagará solo en proporción a la prima pagada.
""")

    st.subheader("Distribución de superficies: bienes comunes vs unidades")
    st.markdown("""
| Fuente | Referencia | Uso |
|--------|------------|-----|
| **Ley 21.442 / Edifito** | Bienes comunes = **50–70% del VR total** | Valor asegurado |
| **ComunidadFeliz** | Bienes comunes = **60–80% del monto asegurado** | Monto póliza |
| **OGUC art. 5.1.11** | Superficie común < **20% sup. útil** no cuenta para constructibilidad | Permisos |
| **Práctica de mercado** | Superficie física común ≈ **30–60% sup. total** según amenidades | Estimación |
| **Reglamento de Copropiedad** | Porcentaje inscrito en Conservador de Bienes Raíces | **Valor legal vinculante** |
""")

    st.subheader("Marco normativo")
    with st.expander("Ley 21.442 — art. 43 (Seguro obligatorio comunidad)"):
        st.markdown("""
Todo condominio habitacional debe contratar seguro colectivo contra incendio cubriendo:
- **Obligatoriamente:** bienes e instalaciones comunes.
- **Opcionalmente:** unidades privadas (el copropietario puede renunciar si tiene póliza propia).
- El copropietario **nunca puede eximirse** del pago por bienes comunes.
""")
    with st.expander("NCG 556 CMF — dic. 2025 (Estructura de la póliza)"):
        st.markdown("""
| Bloque | Cubre | Asegurado | Carácter |
|--------|-------|-----------|----------|
| **1 — Bienes comunes** | Estructura, instalaciones, áreas comunes | La comunidad | Obligatorio |
| **2 — Unidades privadas** | Cada depto, local, bodega | El copropietario | Opcional colectivo |

Ante daños parciales en una unidad, la indemnización se destina **primero a reparación**, no al crédito hipotecario.
""")
    with st.expander("CCom Art. 553 — Regla proporcional"):
        st.markdown("""
Si el monto asegurado < valor real → la compañía paga solo en proporción a la prima.

> Si asegura el **70%** del valor real → recibirá solo el **70%** del daño, aunque el siniestro sea parcial.

**Por eso es fundamental calcular y asegurar el valor correcto.**
""")
    with st.expander("Pasos del cálculo de VR"):
        st.markdown("""
| Paso | Concepto |
|------|----------|
| 1 | VUB (UF/m²) ingresado por el usuario |
| 2 | × Factor geográfico: Metropolitana 1.05 / Intermedia 1.00 / Aislada 1.15 |
| 3 | × Factor normativo: <1985→1.15 / 1985-2000→1.10 / 2001-2010→1.05 / >2010→1.00 |
| 4 | × Factor altura: 1-2p→1.00 / 3-5p→1.05 / 6-10p→1.10 / 11+→1.15 |
| 5 | = Costo directo |
| 6 | + Indirectos 31%: diseño 3% + GG 6% + utilidad 12% + imprevistos 10% |
| 7 | + IVA 19% (si corresponde) |
| ✓ | = Valor de Reconstrucción |
""")
    with st.expander("Fuentes del VUB"):
        st.markdown("""
- **Tabla MINVU** (oficial, en pesos, trimestral): [minvu.gob.cl](https://www.minvu.gob.cl/elementos-tecnicos/tabla-de-costos-unitarios/)
- **Tasador habilitado** — el más preciso para cada caso específico
- **Corredor de seguros** — usa tablas validadas por la compañía aseguradora
- **Referencias de mercado** (orientativas): incluidas en la app al seleccionar zona, tipo y nivel
""")

    st.divider()
    st.caption("Programa referencial. No reemplaza tasación profesional ni asesoría de corredor certificado.")
